"""File scanning and classification tools for the ISA-Tox RO-Crate Builder.

The scanner examines an input directory (or zip archive) and builds a raw
file inventory (path, size, mime type). It is called during session
initialisation, before the agent loop starts, to give the agent a picture
of what files are available.
"""

from __future__ import annotations

import logging
import mimetypes
import os
import shutil
import sys
import tarfile
import time
import zipfile
from collections import Counter
from pathlib import Path

from builder.state import ArchivePreview, FileClassification
from builder.tools.registry import TOOL_REGISTRY

logger = logging.getLogger(__name__)

# Initialise mimetypes so .txt -> text/plain, .csv -> text/csv, etc.
mimetypes.init()

_ZIP_EXTENSIONS = frozenset({".zip", ".tar", ".gz", ".tgz", ".tar.gz"})

# Zip-bomb guard: refuse to extract more than this much *uncompressed* data in
# total from a single archive.
_MAX_UNCOMPRESSED_BYTES = 2 * 1024**3  # 2 GiB


class _UnsafeArchiveError(Exception):
    """An archive member would escape the destination (Zip-Slip) or blow the
    uncompressed-size cap (zip-bomb)."""


def _is_within(dest: Path, target: Path) -> bool:
    """Return True if *target* resolves inside *dest* (Zip-Slip guard)."""
    try:
        target.resolve().relative_to(dest.resolve())
        return True
    except ValueError:
        return False


# Scientific-format MIME registry (Issue #148).
#
# ``mimetypes.guess_type`` returns ``None`` for mass-spec / microscopy / flow
# formats (.mzML, .raw, .wiff, .fcs, .czi, .nd2, .lif, .d, ...), so the old
# text-sniff fallback would decode their bytes and mislabel them as
# ``text/plain``. This map is consulted BEFORE the text sniff so these formats
# get a meaningful media type, and binary vendor formats default to
# ``application/octet-stream`` (the true default for unknown binary content)
# rather than ``text/plain``.
#
# Where a registered/community media type exists we use it (mzML ->
# application/x-mzml, FCS -> application/vnd.isac.fcs); otherwise the format is
# an opaque vendor binary and maps to application/octet-stream.
_SCIENTIFIC_MIME_TYPES: dict[str, str] = {
    # Mass spectrometry — open/standard
    ".mzml": "application/x-mzml",
    ".mzxml": "application/x-mzxml",
    ".mzdata": "application/x-mzdata",
    ".imzml": "application/x-mzml",
    ".mgf": "text/plain",
    # Mass spectrometry — vendor binaries
    ".raw": "application/octet-stream",  # Thermo / Waters raw
    ".wiff": "application/octet-stream",  # SCIEX
    ".wiff2": "application/octet-stream",
    ".d": "application/octet-stream",  # Agilent / Bruker directory bundle
    ".baf": "application/octet-stream",  # Bruker
    ".tdf": "application/octet-stream",  # Bruker timsTOF
    ".yep": "application/octet-stream",  # Bruker
    ".fid": "application/octet-stream",  # Bruker / NMR FID
    # Flow cytometry
    ".fcs": "application/vnd.isac.fcs",
    # Microscopy / imaging — vendor binaries
    ".czi": "application/octet-stream",  # Zeiss
    ".nd2": "application/octet-stream",  # Nikon
    ".lif": "application/octet-stream",  # Leica
    ".lsm": "application/octet-stream",  # Zeiss
    ".oib": "application/octet-stream",  # Olympus
    ".oif": "application/octet-stream",  # Olympus
    ".ims": "application/octet-stream",  # Imaris
    ".vsi": "application/octet-stream",  # Olympus
    ".ndpi": "application/octet-stream",  # Hamamatsu
    ".svs": "application/octet-stream",  # Aperio
    ".scn": "application/octet-stream",  # Leica slide
    ".dm3": "application/octet-stream",  # Gatan
    ".dm4": "application/octet-stream",  # Gatan
    # Open imaging containers
    ".ome.tiff": "image/tiff",
    ".ome.tif": "image/tiff",
}


def _scientific_mime_for(file_path: Path) -> str | None:
    """Return a registered media type for known scientific file extensions.

    Handles both single-suffix (``.mzML``) and compound (``.ome.tiff``)
    extensions, case-insensitively. Returns ``None`` when the extension is not
    in the scientific registry.
    """
    # Prefer the longest matching compound suffix (e.g. .ome.tiff over .tiff).
    suffixes = [s.lower() for s in file_path.suffixes]
    for start in range(len(suffixes)):
        compound = "".join(suffixes[start:])
        if compound in _SCIENTIFIC_MIME_TYPES:
            return _SCIENTIFIC_MIME_TYPES[compound]
    return None


def _detect_mime_type(file_path: Path) -> str:
    """Detect the MIME type of a file using mimetypes and content sniffing.

    Resolution order (Issue #148):

    1. ``mimetypes.guess_type`` (stdlib extension table).
    2. The scientific-format registry (:data:`_SCIENTIFIC_MIME_TYPES`) for
       MS / microscopy / flow extensions the stdlib does not know.
    3. A text-content sniff (CSV / TSV / plain text).
    4. ``application/octet-stream`` as the true default for unknown binary.

    Args:
        file_path: Path to the file.

    Returns:
        A MIME type string.
    """
    # Step 1: stdlib extension table.
    mime_type, _ = mimetypes.guess_type(str(file_path))
    if mime_type:
        return mime_type

    # Step 2: scientific-format registry (BEFORE the text sniff so binary
    # vendor formats are not mislabeled text/plain).
    scientific = _scientific_mime_for(file_path)
    if scientific is not None:
        return scientific

    # Step 3: content sniffing for common text formats.
    try:
        with file_path.open("rb") as f:
            header = f.read(512)

        # A NUL byte is a reliable binary marker — even though it decodes as
        # valid UTF-8, it never appears in genuine text, so refuse to call such
        # content text/* and fall through to octet-stream.
        if b"\x00" not in header:
            # Check for CSV/TSV by looking for commas/tabs in the first line
            if header.startswith(b",") or (b"," in header.split(b"\n")[0][:256]):
                return "text/csv"
            if header.startswith(b"\t") or (b"\t" in header.split(b"\n")[0][:256]):
                return "text/tab-separated-values"

            # Check if it looks like plain text (printable ASCII or common UTF-8)
            try:
                header.decode("utf-8")
                return "text/plain"
            except (UnicodeDecodeError, UnicodeError):
                pass
    except PermissionError:
        logger.warning("Permission denied reading: %s", file_path)
    except OSError:
        pass

    # Step 4: true default for unknown binary content.
    return "application/octet-stream"


def encoding_format_for_name(name: str) -> str | None:
    """Derive an IANA media type from a file name/extension alone (no disk read).

    Used to auto-populate ``schema:encodingFormat`` when drafting File entities
    (Issue #148). Consults the stdlib ``mimetypes`` table first, then the
    scientific-format registry (:data:`_SCIENTIFIC_MIME_TYPES`). Returns
    ``None`` when the name carries no recognised extension, so callers can leave
    ``encodingFormat`` unset rather than guess.

    Args:
        name: A file name or crate-relative path (e.g. ``run.mzML``,
            ``data/results.csv``).

    Returns:
        A media-type string, or ``None`` if the extension is unknown.
    """
    if not name:
        return None
    path = Path(name)
    if not path.suffix:
        return None
    mime_type, _ = mimetypes.guess_type(path.name)
    if mime_type:
        return mime_type
    return _scientific_mime_for(path)


_TABULAR_MIME_TYPES = {
    "text/csv",
    "text/tab-separated-values",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

_TABULAR_SUFFIXES = {".csv", ".tsv", ".xlsx", ".xls"}


# ---------------------------------------------------------------------------
# Archive helpers
# ---------------------------------------------------------------------------


def _is_archive(path: Path) -> bool:
    """Return True if *path* looks like a compressed archive."""
    return path.suffix.lower() in _ZIP_EXTENSIONS or path.name.endswith(".tar.gz")


def _list_archive_contents(archive_path: Path) -> list[dict]:
    """Peek inside a zip *or* tar(.gz) archive and return metadata about contents.

    Returns a list of dicts with keys ``path``, ``size``, ``is_dir`` so the
    agent can present a preview to the user without extracting. Returns ``[]``
    for corrupt or unsupported archives.
    """
    contents: list[dict] = []
    try:
        if zipfile.is_zipfile(archive_path):
            with zipfile.ZipFile(archive_path, "r") as zf:
                for info in zf.infolist():
                    contents.append(
                        {"path": info.filename, "size": info.file_size, "is_dir": info.is_dir()}
                    )
        elif tarfile.is_tarfile(archive_path):
            with tarfile.open(archive_path, "r:*") as tf:
                for member in tf.getmembers():
                    contents.append(
                        {"path": member.name, "size": member.size, "is_dir": member.isdir()}
                    )
        else:
            logger.warning("Not a supported archive (zip/tar): %s", archive_path)
            return []
    except (zipfile.BadZipFile, tarfile.TarError):
        logger.warning("Bad/corrupt archive: %s", archive_path)
        return []
    except Exception:
        logger.exception("Error reading archive: %s", archive_path)
        return []
    contents.sort(key=lambda c: (not c["is_dir"], c["path"]))
    return contents


def preview_archive(path: str) -> ArchivePreview:
    """Return metadata about a zip archive without extracting it.

    Provides a preview of the archive's contents (entry paths, sizes, and
    whether each entry is a directory) so the agent can present this to the
    user before deciding to extract.

    Handles gracefully:
    - Non-existent files: returns an ArchivePreview with an error message
      and an empty entries list.
    - Corrupt/invalid zip files: returns an ArchivePreview with an error
      message and an empty entries list.

    Args:
        path: Path to the zip archive to preview.

    Returns:
        An ``ArchivePreview`` dataclass with archive metadata.
    """
    archive_path = Path(path).resolve()

    # Handle non-existent file
    if not archive_path.is_file():
        return ArchivePreview(
            path=str(archive_path),
            filename=archive_path.name,
            size_bytes=0,
            size_mb=0.0,
            entry_count=0,
            entries=[],
            message=f"File not found: {path}",
            error=f"File not found: {path}",
        )

    size_bytes = archive_path.stat().st_size
    size_mb = size_bytes / (1024 * 1024)

    # Use existing helper — returns [] on corrupt/unsupported archives
    entries = _list_archive_contents(archive_path)

    if not entries and size_bytes > 0:
        # _list_zip_contents returns [] for corrupt zips
        return ArchivePreview(
            path=str(archive_path),
            filename=archive_path.name,
            size_bytes=size_bytes,
            size_mb=round(size_mb, 2),
            entry_count=0,
            entries=[],
            message=f"Cannot read archive: {archive_path.name} (corrupt or unsupported format)",
            error=f"Cannot read archive: {archive_path.name} (corrupt or unsupported format)",
        )

    return ArchivePreview(
        path=str(archive_path),
        filename=archive_path.name,
        size_bytes=size_bytes,
        size_mb=round(size_mb, 2),
        entry_count=len(entries),
        entries=entries,
        message=(f"Archive {archive_path.name}: {size_mb:.1f} MB, {len(entries)} entries"),
    )


def _safe_walk(root: Path) -> list[Path]:
    """Recursively walk *root* skipping unreadable directories.

    ``Path.rglob`` raises ``PermissionError`` on unreadable subdirectories
    (e.g. ``/proc/*/map_files``).  This uses ``os.walk`` with an ``onerror``
    handler that logs and continues rather than crashing.

    Note that ``os.walk`` on Python 3.12+ raises ``PermissionError`` from
    ``os.scandir`` by default, so the ``onerror`` callback is essential.

    Hidden and special directories (``.git``, ``__MACOSX``, dot-prefixed)
    are pruned in-place during the walk so they are never descended.
    """
    results: list[Path] = []

    def _onerror(err: OSError) -> None:
        logger.warning("Skipping unreadable directory: %s", err.filename or err)

    def _should_prune(dirname: str) -> bool:
        """Return True if *dirname* should be skipped entirely."""
        return dirname.startswith(".") or dirname == "__MACOSX"

    try:
        for dirpath_str, dirnames, filenames in os.walk(root, onerror=_onerror):
            # Prune hidden and special directories in-place so os.walk
            # never descends into them (Issue #69).
            dirnames[:] = [d for d in dirnames if not _should_prune(d)]

            dirpath = Path(dirpath_str)
            for name in filenames:
                results.append(dirpath / name)
    except PermissionError:
        # os.walk's onerror catches PermissionError during scandir, but
        # a top-level PermissionError on the root itself still bubbles.
        logger.warning("Permission denied at root: %s", root)
    return results


def scan_files(
    path: str,
    approved_roots: set[str] | None = None,
    max_files: int = 0,
    max_line_length: int = 0,
) -> list[FileClassification]:
    """Scan a directory or zip archive and return a file inventory.

    If *path* is a directory, walks it and returns a list of
    ``FileClassification`` records.  If *path* is a zip archive (``.zip``,
    ``.tar.gz``, etc.), it is automatically extracted to a temporary
    directory and the extracted contents are scanned transparently.

    .. security::

       The scanner **only** accepts paths that descend from one of the
       *approved_roots* — directories the user has explicitly permitted.
       If *approved_roots* is ``None`` (the default) the first path scanned
       is auto-approved and becomes the sole root for subsequent calls.

    Args:
        path: Path to the directory **or** archive to inspect.
        approved_roots: Set of resolved absolute directory paths that are
            allowed for scanning.  Pass ``None`` on the first call to
            auto-approve the target.
        max_files: Maximum number of files to return. 0 means unlimited.
            When exceeded, the list is truncated and a warning is logged.
        max_line_length: Maximum length for each ``first_rows`` line.
            0 means unlimited. Longer lines are truncated.

    Returns:
        A list of ``FileClassification`` records.
    """
    target = Path(path).resolve()

    # -- Security guard: approved-roots check -----------------------------------
    if approved_roots is not None:
        if not any(str(target) == r or str(target).startswith(r + "/") for r in approved_roots):
            logger.warning(
                "Refusing to scan %s — not in approved roots: %s",
                target,
                approved_roots,
            )
            return []

    # -- Security guard: don't follow symlinks out of the resolved path ---------
    _follows_symlinks = False
    try:
        _follows_symlinks = target.is_symlink()
    except PermissionError:
        pass
    if _follows_symlinks:
        logger.warning("Refusing to scan symlink: %s", target)
        return []

    # -- Archive case: auto-extract and recurse --------------------------------
    if target.is_file() and _is_archive(target):
        contents = _list_archive_contents(target)
        size_mb = target.stat().st_size / (1024 * 1024)
        logger.info(
            "Path %s is an archive (%.1f MB) with %d entries — auto-extracting",
            target,
            size_mb,
            len(contents),
        )
        result = unzip_file(str(target))
        if "error" in result:
            logger.warning("Could not extract archive: %s", result["error"])
            return []
        extracted_dir = result["extracted_to"]
        # Scan extracted directory recursively
        return scan_files(extracted_dir)

    # -- Directory case --------------------------------------------------------
    if not target.is_dir():
        logger.warning("Path not found: %s", target)
        return []

    results: list[FileClassification] = []
    try:
        # Use _safe_walk which prunes hidden/special dirs in-place (Issue #69)
        all_entries = sorted(_safe_walk(target))
    except PermissionError:
        logger.warning("Permission denied scanning %s", target)
        return []

    scan_start = time.monotonic()
    processed = 0
    total_candidates = len(all_entries)

    for entry in all_entries:
        if not entry.is_file():
            continue

        # Skip hidden files (names starting with '.')
        if entry.name.startswith("."):
            continue

        # Compute stat and mime once per file (Issue #67)
        try:
            st = entry.stat()
            file_size = st.st_size
        except OSError:
            continue

        mime = _detect_mime_type(entry)

        # Check if this is a tabular file that should have first_rows
        is_tabular = mime in _TABULAR_MIME_TYPES or entry.suffix.lower() in _TABULAR_SUFFIXES
        first_rows: list[str] | None = None
        if is_tabular:
            # For binary office formats (xlsx/xls) the extension says "tabular" but
            # the content is a zip/OLE2 container — do NOT pass already_text=True
            # because that bypasses the NUL-byte binary guard and the raw PK\x03\x04…
            # bytes would be read as mojibake (Issue #51 regression).
            _is_binary_ext = entry.suffix.lower() in {".xlsx", ".xls"}
            sample = read_file_sample(
                str(entry),
                lines=20,
                precomputed_size=file_size,
                already_text=not _is_binary_ext,
            )
            if sample is not None:
                lines = sample.splitlines()
                if max_line_length > 0:
                    lines = [line[:max_line_length] for line in lines]
                first_rows = lines

        results.append(
            FileClassification(
                path=str(entry),
                filename=entry.name,
                size=file_size,
                mime_type=mime,
                first_rows=first_rows,
            )
        )
        processed += 1
        if processed % 100 == 0:
            elapsed = time.monotonic() - scan_start
            logger.debug(
                "Progress: %d/%d files... (%.1fs elapsed)",
                processed,
                total_candidates,
                elapsed,
            )

        # Check max_files cap (Issue #68)
        if max_files > 0 and len(results) >= max_files:
            logger.warning(
                "Reached max_files limit (%d) — truncating scan results",
                max_files,
            )
            break

    total_elapsed = time.monotonic() - scan_start
    logger.info(
        "Scan complete: %d files in %.2fs",
        len(results),
        total_elapsed,
    )
    return results


def _extract_zip_safely(archive_path: Path, dest: Path) -> int:
    """Extract a zip into *dest*, rejecting Zip-Slip members and capping total
    uncompressed size. Returns the entry count. Raises :class:`_UnsafeArchiveError`."""
    with zipfile.ZipFile(archive_path, "r") as zf:
        infos = zf.infolist()
        total = 0
        for info in infos:
            if not _is_within(dest, dest / info.filename):
                raise _UnsafeArchiveError(
                    f"member escapes destination (Zip-Slip): {info.filename!r}"
                )
            total += info.file_size
            if total > _MAX_UNCOMPRESSED_BYTES:
                raise _UnsafeArchiveError(
                    f"uncompressed size exceeds cap ({_MAX_UNCOMPRESSED_BYTES} bytes)"
                )
        zf.extractall(dest)
        return len(infos)


def _extract_tar_safely(archive_path: Path, dest: Path) -> int:
    """Extract a tar(.gz) into *dest* with the same Zip-Slip + size-cap guards,
    plus the 3.12 ``data`` filter for member/link sanitization."""
    with tarfile.open(archive_path, "r:*") as tf:
        members = tf.getmembers()
        total = 0
        for member in members:
            if not _is_within(dest, dest / member.name):
                raise _UnsafeArchiveError(
                    f"member escapes destination (Zip-Slip): {member.name!r}"
                )
            total += member.size
            if total > _MAX_UNCOMPRESSED_BYTES:
                raise _UnsafeArchiveError(
                    f"uncompressed size exceeds cap ({_MAX_UNCOMPRESSED_BYTES} bytes)"
                )
        # filter="data" (Python 3.12+) rejects absolute paths, traversal, and
        # unsafe symlink/hardlink targets that the name check above can't catch.
        tf.extractall(dest, filter="data")
        return len(members)


def unzip_file(path: str, output_dir: str | None = None) -> dict:
    """Safely extract a zip or tar(.gz) archive and return the extracted directory.

    Guards against Zip-Slip (members escaping the destination) and zip-bombs
    (total uncompressed size is capped at ``_MAX_UNCOMPRESSED_BYTES``); unsafe
    archives are refused with a clean error rather than raising. The agent should
    call ``present_to_human`` before calling this tool for large archives so the
    user can confirm.

    Args:
        path: Path to the ``.zip`` / ``.tar`` / ``.tar.gz`` / ``.tgz`` file.
        output_dir: Optional output directory. Defaults to ``<name>_extracted``
            next to the archive.

    Returns:
        A dict with keys ``extracted_to``, ``entry_count``, and ``message`` — or
        ``error``/``message`` on failure.
    """
    archive_path = Path(path)
    if not archive_path.is_file():
        return {
            "error": f"File not found: {path}",
            "message": "The file does not exist.",
        }

    if output_dir:
        dest = Path(output_dir)
        dest.mkdir(parents=True, exist_ok=True)
    else:
        # Extract next to the archive by default — survives session resume
        dest = archive_path.parent / f"{archive_path.stem}_extracted"
        dest.mkdir(parents=True, exist_ok=True)

    try:
        if zipfile.is_zipfile(archive_path):
            extracted = _extract_zip_safely(archive_path, dest)
        elif tarfile.is_tarfile(archive_path):
            extracted = _extract_tar_safely(archive_path, dest)
        else:
            return {
                "error": f"Cannot open {path} as a zip or tar archive (corrupt/unsupported?)",
                "message": "The archive appears to be corrupt or not a supported format.",
            }
    except _UnsafeArchiveError as exc:
        logger.warning("Refused to extract unsafe archive %s: %s", path, exc)
        return {
            "error": f"Unsafe archive refused: {exc}",
            "message": f"Refused to extract {archive_path.name}: {exc}.",
        }
    except (zipfile.BadZipFile, tarfile.TarError) as exc:
        return {
            "error": f"Cannot extract {path} (corrupt?): {exc}",
            "message": "The archive appears to be corrupt or not a valid zip/tar.",
        }
    except Exception as exc:
        logger.exception("Error extracting archive")
        return {"error": str(exc), "message": f"Failed to extract: {exc}"}

    # Strip macOS resource-fork metadata (__MACOSX, ._ files) unless we're on macOS
    _strip_macos_junk(dest)

    logger.info("Extracted %s -> %s (%d entries)", path, dest, extracted)
    return {
        "extracted_to": str(dest),
        "entry_count": extracted,
        "message": (
            f"Extracted {extracted} entries to {dest}. "
            f"Use ``scan_files`` with path ``{dest}`` to inventory the contents."
        ),
    }


def _strip_macos_junk(root: Path) -> None:
    """Remove __MACOSX directories and ._ resource fork files from *root*.

    These are macOS-specific metadata that get bundled into zips when
    created on a Mac.  On Linux they are useless bloat; on macOS we
    leave them alone since they may be meaningful to the system.
    """
    if sys.platform == "darwin":
        return  # macOS — leave them be
    if not root.is_dir():
        return
    for entry in list(root.rglob("*")):
        # Remove __MACOSX directories and any ._ resource fork files
        if "__MACOSX" in entry.parts:
            try:
                if entry.is_dir():
                    shutil.rmtree(entry)
                else:
                    entry.unlink()
                logger.debug("Removed macOS junk: %s", entry)
            except OSError:
                logger.warning("Could not remove macOS junk: %s", entry)


# ---------------------------------------------------------------------------
# Summary helpers (used by read_file_sample with mode="summary")
# ---------------------------------------------------------------------------

_SUMMARY_MAX_LINES = 50


def _looks_binary(path: Path) -> bool:
    """Return True if the file has a NUL byte in its first 8 KiB.

    A reliable heuristic for binary content. Office formats (.xls/.xlsx) and
    GraphPad files (.prism/.pzf) are binary containers that ``mimetypes`` may
    mislabel as ``text/csv``; this guard stops them being read as mojibake text.
    """
    try:
        with path.open("rb") as fb:
            return b"\x00" in fb.read(8192)
    except OSError:
        return True


def _summarize_csv(path: Path) -> str | None:
    """Return a summary of a CSV/TSV file: columns, row count, sample."""
    try:
        with path.open("r", encoding="utf-8", errors="replace") as f:
            header = f.readline().strip()
            if not header:
                return None
            columns = [c.strip() for c in header.replace("\t", ",").split(",")]
            n_cols = len(columns)
            data_rows = 0
            sample_rows: list[str] = []
            for line in f:
                stripped = line.rstrip("\n")
                data_rows += 1
                if len(sample_rows) < 3:
                    sample_rows.append(stripped)
    except OSError:
        return None

    parts = [
        "Format: CSV/TSV",
        f"Columns ({n_cols}): {', '.join(columns)}",
        f"Data rows: {data_rows}",
    ]
    if sample_rows:
        parts.append("Sample rows:")
        parts.extend(f"  {r}" for r in sample_rows)
    return "\n".join(parts)


def _summarize_json(path: Path) -> str | None:
    """Return a summary of a JSON file: keys, array lengths, structure."""
    import json as json_lib

    try:
        with path.open("r", encoding="utf-8", errors="replace") as f:
            raw = f.read(200_000)
        data = json_lib.loads(raw)
    except (json_lib.JSONDecodeError, OSError, ValueError):
        return None

    parts = ["Format: JSON"]

    if isinstance(data, dict):
        keys = list(data.keys())
        parts.append(f"Top-level keys ({len(keys)}): {', '.join(str(k) for k in keys)}")
        for k in keys:
            v = data[k]
            if isinstance(v, list):
                parts.append(f"  {k}: array[{len(v)}]")
            elif isinstance(v, dict):
                subkeys = list(v.keys())
                parts.append(f"  {k}: object {{{', '.join(str(sk) for sk in subkeys[:5])}}}")
            elif isinstance(v, str):
                preview = v[:80].replace("\n", " ")
                parts.append(f'  {k}: string "{preview}"')
            else:
                parts.append(f"  {k}: {type(v).__name__} = {v}")
    elif isinstance(data, list):
        parts.append(f"Top-level array with {len(data)} items")
        if data and isinstance(data[0], dict):
            keys = list(data[0].keys())
            parts.append(f"  Item keys: {', '.join(str(k) for k in keys)}")
    else:
        parts.append(f"Top-level value: {type(data).__name__} = {data}")

    return "\n".join(parts)


def _summarize_xlsx(path: Path) -> str | None:
    """Summary of XLSX: sheet names and row/col counts."""
    try:
        import openpyxl
    except ImportError:
        return None

    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True, keep_links=False)
    except Exception:
        return None

    parts = ["Format: Excel (.xlsx)"]
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # read_only worksheets have no `.dimensions`; count by iterating
            # (capped, so a huge sheet doesn't make the summary slow).
            header: tuple = ()
            col_count = 0
            row_count = 0
            capped = False
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i == 0:
                    header = row
                    col_count = len(row)
                row_count = i + 1
                if i >= 2000:
                    capped = True
                    break
            cols = ", ".join(str(c) for c in header if c is not None)
            rows_label = f"{row_count}+" if capped else str(row_count)
            line = f"  {sheet_name}: {rows_label} rows x {col_count} cols"
            if cols:
                line += f"; columns: {cols}"
            parts.append(line)
    except Exception:
        # A malformed sheet must not crash (or escape) the summary.
        if len(parts) == 1:
            return None
    finally:
        wb.close()

    return "\n".join(parts)


def _summarize_pdf(path: Path) -> str | None:
    """Summary of PDF: page count, text length."""
    # Try pdfplumber first
    try:
        import pdfplumber
    except ImportError:
        pdfplumber = None  # ty: ignore[invalid-assignment]

    if pdfplumber is not None:
        try:
            with pdfplumber.open(path) as pdf:
                try:
                    n_pages = len(pdf.pages)
                except Exception:
                    n_pages = 0
                total_chars = 0
                for page in pdf.pages:
                    try:
                        total_chars += len(page.extract_text() or "")
                    except Exception:
                        pass
                return (
                    f"Format: PDF\n"
                    f"Pages: {n_pages}\n"
                    f"Total text length: ~{total_chars} characters"
                )
        except Exception:
            pass  # Fall through to the manual scan below

    # Fallback: count /Type /Page entries manually
    try:
        with path.open("rb") as f:
            content = f.read(10000)
        page_count = content.count(b"/Type /Page") or content.count(b"/Type/Page")
        if page_count:
            return f"Format: PDF\nPages: ~{page_count}"
    except OSError:
        pass
    return None


def _summarize_docx(path: Path) -> str | None:
    """Summary of DOCX: paragraph count, table count, sample."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        return None

    try:
        doc = Document(str(path))
    except Exception:
        return None

    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    total_chars = sum(len(p.text) for p in doc.paragraphs)

    parts = [
        "Format: Word (.docx)",
        f"Paragraphs: {para_count}",
        f"Tables: {table_count}",
        f"Total text length: ~{total_chars} characters",
    ]

    sample_paras = []
    for p in doc.paragraphs:
        if p.text.strip():
            sample_paras.append(p.text.strip()[:120])
            if len(sample_paras) >= 3:
                break
    if sample_paras:
        parts.append("Sample text:")
        parts.extend(f"  {s}" for s in sample_paras)

    return "\n".join(parts)


def _summarize_text(path: Path) -> str | None:
    """Summary of plain text: line count, char count, sample."""
    try:
        file_size = path.stat().st_size
    except OSError:
        return None

    line_count = 0
    sample_lines: list[str] = []
    try:
        with path.open("r", encoding="utf-8", errors="replace") as f:
            for line in f:
                line_count += 1
                if len(sample_lines) < 5:
                    sample_lines.append(line.rstrip("\n"))
    except OSError:
        return None

    parts = [
        "Format: Plain text",
        f"Lines: {line_count}",
        f"Size: {file_size} bytes",
    ]
    if sample_lines:
        parts.append("Sample:")
        parts.extend(f"  {s}" for s in sample_lines)
    return "\n".join(parts)


def _build_file_overview(path: str, file_path: Path, summary: str | None) -> str:
    """Build an overview string with file metadata plus the summary."""
    file_size = 0
    mime_type = "unknown"
    try:
        file_size = file_path.stat().st_size
        mime_type = _detect_mime_type(file_path)
    except OSError:
        pass

    parts = [
        f"File: {path}",
        f"Size: {_format_size(file_size)}",
        f"MIME type: {mime_type}",
    ]
    if summary:
        parts.append("")
        parts.append(summary)
    return "\n".join(parts)


def _format_size(size_bytes: int) -> str:
    """Format bytes as human-readable string."""
    if size_bytes == 0:
        return "0 B"
    size = float(size_bytes)
    for unit in ("B", "KB", "MB", "GB"):
        if abs(size) < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"


def read_multiple_files(
    paths: list[str],
    *,
    lines: int = 50,
    mode: str = "content",
) -> dict:
    """Read several files in one go and return their contents.

    Use this tool when you need to inspect multiple files at once — for
    example all the metadata files in an assay directory — rather than
    calling ``read_file_sample`` for each individually.

    Args:
        paths: List of file paths to read (absolute or relative to cwd).
        lines: Max lines to read per file (default 50).
        mode: Reading mode - "content" (first N lines), "summary" (file-type-aware
            summary), or "overview" (metadata + summary). Default "content".

    Returns:
        A dict with:
        - ``files``: ``{path: content_or_error}``
        - ``count``: number of files successfully read
        - ``skipped``: list of paths that could not be read
    """
    results: dict[str, str] = {}
    skipped: list[str] = []
    total_start = time.monotonic()

    for path in paths:
        file_start = time.monotonic()
        content = read_file_sample(path, lines=lines, mode=mode)
        file_elapsed = time.monotonic() - file_start
        if content is not None:
            results[path] = content
            logger.debug("Read %s in %.3fs", path, file_elapsed)
        else:
            skipped.append(path)
            logger.debug("Skipped %s in %.3fs", path, file_elapsed)

    total_elapsed = time.monotonic() - total_start
    count = len(results)
    logger.info(
        "Read %d file(s) from %d path(s) in %.2fs",
        count,
        len(paths),
        total_elapsed,
    )
    msg = f"Read {count} file(s)"
    if skipped:
        msg += f" ({len(skipped)} skipped: {', '.join(skipped)})"
    return {
        "files": results,
        "count": count,
        "skipped": skipped,
        "message": msg,
    }


def read_file_sample(
    path: str,
    lines: int = 20,
    precomputed_size: int | None = None,
    already_text: bool = False,
    mode: str = "content",
) -> str | None:
    """Read a sample from a file with optional mode-aware output.

    Args:
        path: Path to the file to sample.
        lines: Number of lines to read (default 20). Used in "content" mode.
        precomputed_size: Precomputed file size to avoid a stat() call.
            When given, the size check uses this value instead of calling stat().
        already_text: If True, skip MIME and binary-content detection because
            the caller has already determined this is a text file. Avoids
            redundant _detect_mime_type and content checks during scanning.
        mode: Reading mode:
            - "content" (default): First N lines of the file (legacy behavior).
            - "summary": File-type-aware summary (columns for CSV, keys for JSON,
              sheet names for XLSX, page count for PDF, etc.).
            - "overview": File metadata (name, size, MIME type) plus the summary.

    Returns:
        A string with the file content/summary, or None if the file cannot be read.
    """
    file_path = Path(path)
    if not file_path.is_file():
        return None

    # For summary and overview modes, dispatch to type-aware summarizers
    if mode in ("summary", "overview"):
        # Determine MIME type for dispatch
        try:
            mime = _detect_mime_type(file_path)
        except Exception:
            mime = "application/octet-stream"

        suffix = file_path.suffix.lower()

        # Pick the right summarizer based on file type
        summary: str | None = None

        # Dispatch by suffix to the binary-aware summarizers FIRST: office
        # formats carry a tabular MIME, so the CSV branch below would otherwise
        # claim them and read their zip/OLE2 header as mojibake "columns".
        if suffix == ".json":
            summary = _summarize_json(file_path)
        elif suffix == ".xlsx" or suffix == ".xls":
            summary = _summarize_xlsx(file_path)
        elif suffix == ".pdf":
            summary = _summarize_pdf(file_path)
        elif suffix == ".docx":
            summary = _summarize_docx(file_path)
        elif suffix == ".csv" or suffix == ".tsv" or mime in _TABULAR_MIME_TYPES:
            # A textual MIME (text/csv, ms-excel) can mislabel a binary file —
            # e.g. GraphPad .prism/.pzf — which _summarize_csv would read as
            # mojibake "Columns: PK\x03\x04…". Refuse binary content here.
            if _looks_binary(file_path):
                return None
            summary = _summarize_csv(file_path)
        else:
            # Try as plain text first, else binary check
            if _looks_binary(file_path):
                return None  # binary — no summary possible
            summary = _summarize_text(file_path)

        if mode == "overview":
            return _build_file_overview(path, file_path, summary)
        return summary  # mode == "summary"

    # Legacy "content" mode: first N lines
    # Skip very large files
    try:
        size = precomputed_size if precomputed_size is not None else file_path.stat().st_size
        if size > 100 * 1024 * 1024:
            logger.info("Skipping file sample (>100MB): %s", path)
            return None
    except OSError:
        return None

    if not already_text:
        # Skip binary files
        if _detect_mime_type(file_path) == "application/octet-stream":
            return None

        # Content-based binary guard: office formats (xlsx/xls) are zip/OLE2
        # containers whose MIME type is *not* octet-stream, so the check above
        # misses them and their raw bytes (PK\x03\x04…) would be read as mojibake.
        # A NUL byte in the first chunk reliably marks a file as binary.
        try:
            with file_path.open("rb") as fb:
                if b"\x00" in fb.read(8192):
                    return None
        except OSError:
            return None

    try:
        sample_lines: list[str] = []
        with file_path.open("r", encoding="utf-8", errors="replace") as f:
            for _ in range(lines):
                line = f.readline()
                if not line:
                    break
                sample_lines.append(line.rstrip("\n"))
            return "\n".join(sample_lines)
    except PermissionError:
        logger.warning("Permission denied reading file: %s", path)
        return None
    except Exception:
        logger.exception("Error reading file sample: %s", path)
        return None


def extract_pdf_text(path: str) -> str | None:
    """Extract structured content from a PDF file using pdfplumber.

    Returns text, tables, and image metadata in a structured format with
    ``[Page N]``, ``[Text]``, ``[Table N]``, and ``[Image]`` section
    markers so the LLM can understand the layout.  Tables are formatted
    as markdown-like pipe-delimited rows.

    Args:
        path: Path to the PDF file.

    Returns:
        A structured string with page/section markers, or None if the file
        cannot be read, is not a valid PDF, is password-protected, or
        exceeds 100 MB.
    """
    file_path = Path(path)
    if not file_path.is_file():
        return None

    # Skip very large files
    try:
        size = file_path.stat().st_size
        if size > 100 * 1024 * 1024:
            logger.info("Skipping large PDF (>100MB): %s", path)
            return None
    except OSError:
        return None

    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber is not installed. Install it with: uv add pdfplumber")
        return None

    try:
        with pdfplumber.open(file_path) as pdf:
            if len(pdf.pages) == 0:
                return ""

            output_parts: list[str] = []

            for page_idx, page in enumerate(pdf.pages, start=1):
                page_parts: list[str] = []
                page_parts.append(f"[Page {page_idx}]")

                # --- Extract text ---
                text = page.extract_text(layout=True)
                if text and text.strip():
                    for line in text.strip().split("\n"):
                        line = line.strip()
                        if line:
                            # pdfplumber encodes newlines as (cid:10) for some
                            # fonts; replace them and any other CID sequences.
                            line = line.replace("(cid:10)", " ")
                            line = line.replace("(cid:13)", " ")
                            line = line.strip()
                            if line:
                                page_parts.append(f"[Text] {line}")

                # --- Extract tables ---
                tables = page.extract_tables()
                if tables:
                    for table_idx, table_data in enumerate(tables, start=1):
                        if not table_data:
                            continue
                        # Clean up: collapse None cells, strip whitespace
                        cleaned_rows: list[list[str]] = []
                        for row in table_data:
                            cleaned = [(cell or "").strip() for cell in row]
                            if any(cleaned):
                                cleaned_rows.append(cleaned)

                        if not cleaned_rows:
                            continue

                        page_parts.append(f"[Table {table_idx} ({len(cleaned_rows)} rows)]")
                        # Headers
                        header = cleaned_rows[0]
                        page_parts.append("| " + " | ".join(header) + " |")
                        # Separator
                        page_parts.append("| " + " | ".join("---" for _ in header) + " |")
                        # Data rows
                        for row in cleaned_rows[1:]:
                            page_parts.append("| " + " | ".join(row) + " |")

                # --- Report images ---
                images = page.images
                if images:
                    img_info: list[str] = []
                    for img in images:
                        w = img.get("width", 0)
                        h = img.get("height", 0)
                        img_info.append(f"  - Image ({w:.0f}x{h:.0f} pts)")
                    if img_info:
                        page_parts.append(f"[Image] {len(images)} image(s) on this page")
                        page_parts.extend(img_info)

                if page_parts:
                    output_parts.append("\n".join(page_parts))

            return "\n\n".join(output_parts)

    except PermissionError:
        logger.warning("Permission denied reading file: %s", path)
        return None
    except Exception:
        logger.exception("Error extracting PDF content from: %s", path)
        return None


def summarize_scan_result(files: list[FileClassification], sample: int = 15) -> str:
    """Return a compact, LLM-facing summary of a scan result.

    The full inventory is stored in ``CrateState.scanned_files``; the agent
    only needs a clear success signal plus a small sample — not the raw list of
    hundreds/thousands of dataclass objects, which floods the context with no
    obvious success cue and makes the agent re-scan in a loop.

    Args:
        files: The scanned ``FileClassification`` records.
        sample: Maximum number of filenames to include in the sample.

    Returns:
        A short human/LLM-readable summary string.
    """
    n = len(files)
    if n == 0:
        return "Scan complete: 0 files found."

    by_type = Counter(f.mime_type for f in files)
    types = ", ".join(f"{mime} ({count})" for mime, count in by_type.most_common(8))
    shown = ", ".join(f.filename for f in files[:sample])
    more = f", +{n - sample} more" if n > sample else ""
    return (
        f"Scan complete: {n} files found and stored in session state "
        f"(no need to scan again). Types: {types}. Sample: {shown}{more}."
    )


# ---------------------------------------------------------------------------
# Explicit ToolRegistry registration
# ---------------------------------------------------------------------------

TOOL_REGISTRY.register(
    "extract_pdf_text",
    extract_pdf_text,
    description="Extract structured content from a PDF file (text, tables, and image metadata)",
)
