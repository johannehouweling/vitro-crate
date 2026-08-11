"""Tests for the deterministic pipeline spine (Issue #179, task 4).

``run_pipeline`` is a pure, code-driven orchestrator over an already-initialized
:class:`~builder.engine.AgentEngine`: it scaffolds the ISA backbone, drafts what
state already has, builds + validates in memory, and runs a bounded deterministic
fix loop — with NO LLM deciding control flow. These tests are fully offline (the
validator runs against the bundled RO-Crate context; see
``tests/test_offline_validation.py``).

The headline guarantees under test:

* the spine on an *empty* state reaches ``{base, isa, tox}`` REQUIRED conformance;
* the bounded fix loop clears a seeded REQUIRED issue;
* the pipeline is **deterministic** — same input ⇒ identical built ``@graph`` hash
  across independent runs.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import pytest

from builder.engine import AgentEngine
from builder.state import (
    CrateState,
    Entity,
    EntityProvenance,
    EntityType,
    FileClassification,
)
from builder.tools.hitl import SimulatedHumanInterface

# The spine drives build_and_validate / fix_required_issues, which run the
# (deliberately uncached, owlrl-heavy) SHACL validator several times — slower than
# the suite-wide CI default. Mirror tests/test_tools_repair.py and give this
# validation-heavy module headroom; the marker overrides the CLI --timeout.
pytestmark = pytest.mark.timeout(120)


@pytest.fixture(autouse=True)
def _isolate_session_dir(tmp_path_factory, monkeypatch):
    """Redirect save_session's SESSION_DIR to a tmp dir for every test here.

    ``run_pipeline`` now persists CrateState at each phase boundary (#242). Point
    the writer at a throwaway dir (and reset its module-level dedup cache) so the
    spine's saves never litter the repo's real ``sessions/`` directory and tests
    stay hermetic.
    """
    import builder.tools.session as sess_mod

    monkeypatch.setattr(sess_mod, "SESSION_DIR", tmp_path_factory.mktemp("sessions"))
    monkeypatch.setattr(sess_mod, "_last_saved_state_hash", None)


def _entity(entity_id: str, type_: EntityType, **fields) -> Entity:
    return Entity(
        entity_id=entity_id,
        type=type_,
        fields=dict(fields),
        _provenance=EntityProvenance(created_by="llm"),
    )


def _engine(state: CrateState | None = None) -> AgentEngine:
    """A headless engine on *state* (no scan), session_id assigned via initialize()."""
    engine = AgentEngine(state=state or CrateState(), human_interface=SimulatedHumanInterface())
    engine.initialize()  # assigns session_id + opens profiler; no input_path => no scan
    return engine


class TestRunPipelineShape:
    def test_returns_result_dict_with_conformance(self) -> None:
        from builder.agents.pipeline.pipeline import run_pipeline

        result = run_pipeline(_engine())
        assert isinstance(result, dict)
        # The result reports the final per-layer conformance map.
        assert "conformance" in result
        assert set(result["conformance"]) == {"base", "isa", "tox"}
        assert "ok" in result


class TestEmptyStateReachesConformance:
    def test_scaffold_only_reaches_base_and_isa(self) -> None:
        """An empty state, run through the spine, becomes {base,isa}-conformant."""
        from builder.agents.pipeline.pipeline import run_pipeline

        engine = _engine()
        result = run_pipeline(engine)

        # BASE and ISA conform from the scaffold alone. TOX does not, and that is
        # correct: an empty state has no exposure duration or detection instrument
        # to state, and `_pv` refuses to publish "unknown" as if it were a
        # measurement (D5). The shape firing IS the prompt to go and supply it.
        conformance = result["conformance"]
        assert conformance == {"base": True, "isa": True, "tox": False}
        assert result["issues"], "tox must report why it did not conform"
        assert all(
            str(issue.get("property", "")).endswith("additionalProperty")
            for issue in result["issues"]
        ), result["issues"]

        # The backbone really exists in state (scaffold step ran via the engine).
        types = {e.type for e in engine.state.list_entities()}
        assert {"Investigation", "Study", "Assay"} <= types

    def test_scaffold_step_supplies_required_study_name(self) -> None:
        """Regression: a bare draft_study has no `name` and fails ISA; the spine
        must supply backbone names deterministically so ISA passes with no LLM."""
        from builder.agents.pipeline.pipeline import run_pipeline

        engine = _engine()
        run_pipeline(engine)
        study = next(e for e in engine.state.list_entities() if e.type == "Study")
        assert study.fields.get("name")  # non-empty


class TestBoundedFixLoop:
    def _backbone(self) -> CrateState:
        state = CrateState()
        state.metadata.title = "Pipeline fix-loop crate"
        state.add_entity(
            _entity("inv1", "Investigation", name="Inv", description="d", identifier="INV-1")
        )
        state.add_entity(
            _entity("st1", "Study", name="St", description="d", investigation_id="inv1")
        )
        state.add_entity(_entity("as1", "Assay", name="As", description="d", study_id="st1"))
        return state

    def test_fix_loop_clears_seeded_required_issue(self) -> None:
        """An EndpointReadout missing its result + exactly one un-wired File is a
        REQUIRED tox issue the deterministic fix loop must clear (link)."""
        from builder.agents.pipeline.pipeline import run_pipeline

        state = self._backbone()
        state.add_entity(
            _entity(
                "er1", "LabProcess", process_type="EndpointReadout", name="Readout", assay_id="as1"
            )
        )
        state.add_entity(_entity("f0", "File", name="raw0.csv", dest_path="data/raw0.csv"))

        engine = _engine(state)
        result = run_pipeline(engine)

        # The seeded issue cleared: the File is now wired. Tox still reports the
        # scaffold's unparameterised chain steps — with no provider and no source
        # documents there is nothing to assert about them, and `_pv` will not
        # publish "unknown" to paper over it (D5). What matters here is that the
        # RESULT issue is gone; anything left must be the parameter gap.
        assert result["conformance"]["base"] is True
        assert result["conformance"]["isa"] is True
        assert all(
            str(issue.get("property", "")).endswith("additionalProperty")
            for issue in result.get("issues") or []
        ), result.get("issues")
        readout = engine.state.get_entity("er1")
        assert readout is not None
        wired = str(readout.fields.get("result") or readout.fields.get("output") or "")
        assert "f0" in wired

    def test_fix_loop_is_bounded(self) -> None:
        """The fix loop reports a bounded round count and never spins forever."""
        from builder.agents.pipeline.pipeline import run_pipeline

        result = run_pipeline(_engine())
        assert isinstance(result.get("fix_rounds"), int)
        assert 0 <= result["fix_rounds"] <= 3


class TestDraftEntitiesWiring:
    """Step 2 (`_draft_entities`) wires the bounded drafter-leaf into the spine.

    The leaf (`draft_entity_fields`) is STUBBED here — no model, no network. The
    contract under test:

    * with a provider configured AND usable context, the leaf's NON-IDENTIFIER
      fields are applied to the relevant state entities;
    * D5 — identifier / `@id` / `entity_id` fields are NEVER set or overwritten;
    * missing fields are filled, existing fields are NOT overwritten;
    * with NO provider configured, the step is a STRICT no-op (nothing mutated);
    * with a provider but NO usable context, the step is also a strict no-op.
    """

    def _seeded_state(self) -> CrateState:
        """A titled crate with a backbone + a bare seeded compound to enrich."""
        state = CrateState()
        state.metadata.title = "TPO inhibition dose-response screen"
        state.metadata.description = "A cell-based in vitro TPO inhibition assay."
        state.add_entity(_entity("inv1", "Investigation", name="Inv"))
        state.add_entity(_entity("st1", "Study", name="St", investigation_id="inv1"))
        state.add_entity(_entity("as1", "Assay", name="As", study_id="st1"))
        # A bare MolecularEntity (only a name) for the leaf to enrich.
        state.add_entity(_entity("chem1", "MolecularEntity", name="Methimazole"))
        return state

    def _enable_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Make `get_provider()` report a configured provider (no real model)."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")

    def test_stub_leaf_applies_non_identifier_fields(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        calls: list[tuple[str, str]] = []

        def fake_leaf(entity_type, context, *, overrides=None, usage_sink=None):
            calls.append((entity_type, context))
            # A descriptive field plus an identifier the leaf would normally strip
            # — assert the wiring NEVER applies the identifier even if present.
            return {"description": f"drafted {entity_type}", "identifier": "FAKE-ID"}

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", fake_leaf)

        engine = _engine(self._seeded_state())
        result = pipeline_mod._draft_entities(engine)

        # The leaf saw real context (the title appears in the gathered context).
        assert calls, "the leaf must be invoked when provider + context are present"
        assert any("TPO inhibition" in ctx for _, ctx in calls)

        # The descriptive field landed on the backbone entities.
        study = engine.state.get_entity("st1")
        assert study is not None
        assert study.fields.get("description") == "drafted Study"

        # The seeded compound is NOT enriched: `MolecularEntity`'s D5-pruned
        # schema exposes only `name`, which it already has, so the leaf is never
        # called for it (#423). This stub returns a `description` that the real
        # leaf could not — see TestDrafterLeafIsSkippedWhenItCannotHelp.
        chem = engine.state.get_entity("chem1")
        assert chem is not None
        assert not chem.fields.get("description")
        assert "MolecularEntity" not in [t for t, _ in calls]

        # D5: the identifier the leaf returned was NOT applied to any entity.
        for ent_id in ("inv1", "st1", "as1", "chem1"):
            ent = engine.state.get_entity(ent_id)
            assert ent is not None
            assert ent.fields.get("identifier") != "FAKE-ID"
            # The entity_id / @id is untouched.
            assert ent.entity_id == ent_id

        # The result is informative.
        assert isinstance(result.get("drafted"), list)
        assert "st1" in result["drafted"]
        assert isinstance(result.get("fields_applied"), int)
        assert result["fields_applied"] >= 1

    def test_does_not_overwrite_existing_fields(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)

        def fake_leaf(entity_type, context, *, overrides=None, usage_sink=None):
            return {"name": "LEAF NAME", "description": "leaf desc"}

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", fake_leaf)

        engine = _engine(self._seeded_state())
        pipeline_mod._draft_entities(engine)

        study = engine.state.get_entity("st1")
        assert study is not None
        # `name` was already set ("St") — the leaf must NOT clobber it…
        assert study.fields.get("name") == "St"
        # …but the missing `description` is filled.
        assert study.fields.get("description") == "leaf desc"

    def test_no_provider_is_strict_noop(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        # No provider configured.
        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: None)

        # The leaf must never be called when there is no provider.
        def boom(*args, **kwargs):  # pragma: no cover - must not run
            raise AssertionError("draft_entity_fields must not run without a provider")

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", boom)

        engine = _engine(self._seeded_state())
        before = {e.entity_id: dict(e.fields) for e in engine.state.list_entities()}

        result = pipeline_mod._draft_entities(engine)

        after = {e.entity_id: dict(e.fields) for e in engine.state.list_entities()}
        assert before == after, "no-provider _draft_entities must mutate nothing"
        assert result.get("drafted") == []
        assert result.get("fields_applied") == 0

    def test_no_context_is_strict_noop(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)

        def boom(*args, **kwargs):  # pragma: no cover - must not run
            raise AssertionError("draft_entity_fields must not run without context")

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", boom)

        # An untitled, undescribed, unscanned crate carries no usable context.
        state = CrateState()
        state.add_entity(_entity("inv1", "Investigation", name="Inv"))
        engine = _engine(state)
        before = {e.entity_id: dict(e.fields) for e in engine.state.list_entities()}

        result = pipeline_mod._draft_entities(engine)

        after = {e.entity_id: dict(e.fields) for e in engine.state.list_entities()}
        assert before == after, "no-context _draft_entities must mutate nothing"
        assert result.get("drafted") == []
        assert result.get("fields_applied") == 0


class TestDrafterLeafIsSkippedWhenItCannotHelp:
    """The leaf is not called when nothing it can return is applicable (#423).

    The model is bound to a D5-pruned structured-output schema
    (:func:`builder.agents.pipeline.leaves._structured_output_schema`), and the
    spine applies only `_DESCRIPTIVE_APPLY_FIELDS` that the entity is *missing*.
    When those two sets do not intersect, the call cannot change state whatever
    the model returns — so it must not be made at all.

    This is not a heuristic about which types are "worth" drafting: it is a
    property of the schema. `MolecularEntity` exposes only `name`, so a compound
    that already has a name (every compound resolved via `resolve_compound`) can
    only ever be missing `description` — a field the model is never offered.
    """

    def _enable_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")

    def _titled_state(self) -> CrateState:
        state = CrateState()
        state.metadata.title = "TPO inhibition dose-response screen"
        state.metadata.description = "A cell-based in vitro TPO inhibition assay."
        return state

    def _record_calls(
        self, monkeypatch: pytest.MonkeyPatch, returns: dict | None = None
    ) -> list[str]:
        """Stub the leaf; return the list it appends each called entity_type to."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        seen: list[str] = []

        def fake_leaf(entity_type, context, *, overrides=None, usage_sink=None):
            seen.append(entity_type)
            return dict(returns or {})

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", fake_leaf)
        return seen

    def test_named_compound_never_reaches_the_leaf(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """A compound with a name is missing only `description`, which the
        MolecularEntity schema does not expose — so no call is made."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        seen = self._record_calls(monkeypatch, {"description": "drafted"})

        state = self._titled_state()
        state.add_entity(_entity("chem1", "MolecularEntity", name="Methimazole"))
        engine = _engine(state)

        result = pipeline_mod._draft_entities(engine)

        assert "MolecularEntity" not in seen, (
            "a named compound cannot be enriched by the leaf — the pruned schema "
            f"offers only 'name', which is already set; calls made: {seen}"
        )
        assert result["fields_applied"] == 0
        chem = engine.state.get_entity("chem1")
        assert chem is not None and not chem.fields.get("description")

    def test_an_unnamed_compound_still_reaches_the_leaf(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """The guard is an intersection, not a per-type blacklist: a compound
        missing its `name` CAN be helped, so the call must still happen."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        seen = self._record_calls(monkeypatch, {"name": "Methimazole"})

        state = self._titled_state()
        state.add_entity(_entity("chem1", "MolecularEntity"))
        engine = _engine(state)

        pipeline_mod._draft_entities(engine)

        assert seen.count("MolecularEntity") == 1
        chem = engine.state.get_entity("chem1")
        assert chem is not None and chem.fields.get("name") == "Methimazole"

    def test_a_described_type_is_unaffected(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Control — `Assay` DOES expose `description`, so the guard must not
        touch it. Without this, "skip everything" would pass the test above."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        seen = self._record_calls(monkeypatch, {"description": "drafted assay"})

        state = self._titled_state()
        state.add_entity(_entity("as1", "Assay", name="As"))
        engine = _engine(state)

        pipeline_mod._draft_entities(engine)

        assert seen.count("Assay") == 1
        assay = engine.state.get_entity("as1")
        assert assay is not None
        assert assay.fields.get("description") == "drafted assay"

    def test_the_guard_matches_the_schema_the_model_is_actually_bound_to(self) -> None:
        """The skip set is derived from the real leaf schema, not hardcoded.

        Pins the property that makes the guard correct: for every draftable type,
        the fields the spine can apply are exactly the visible-schema fields it
        intersects. If someone adds `description` to the MolecularEntity draft
        schema, the guard must start calling the leaf again on its own.
        """
        pytest.importorskip("langchain_core")
        from builder.agents.pipeline.leaves import _structured_output_schema

        visible = set(_structured_output_schema("MolecularEntity").get("properties", {}))
        assert "description" not in visible, (
            "this test encodes WHY named compounds are skipped; if MolecularEntity "
            "gains a description field the skip is no longer correct"
        )
        assert "description" in set(_structured_output_schema("Assay").get("properties", {})), (
            "Assay must still expose description, else the control test is vacuous"
        )


class TestDrafterPromptCarriesEntityIdentity:
    """Each drafted entity gets its own context, not one shared digest (#423).

    `draft_entity_fields` receives only `(entity_type, context)`. When `context`
    is the same crate-wide digest for every entity, two entities of one type send
    a byte-identical prompt and the model returns the same text twice — observed
    on a real build as the parental line `cell_cho_k1` being described as
    "stably transfected with ... OATP1C1", copied from its transfected sibling.

    The fix is the pattern the sibling caller already uses
    (`guidance.py::_draft_context`): fold the entity's own identity into the
    free-text context. No signature change.
    """

    def _enable_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")

    def test_two_same_type_entities_get_different_prompts(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        contexts: dict[str, str] = {}
        order: list[str] = []

        def fake_leaf(entity_type, context, *, overrides=None, usage_sink=None):
            order.append(context)
            return {"description": f"desc for context #{len(order)}"}

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", fake_leaf)

        state = CrateState()
        state.metadata.title = "OATP1C1 T4 uptake inhibition screen"
        state.metadata.description = "Transfected CHO-K1 cell model."
        state.add_entity(_entity("cell_cho_k1", "CellLineSample", name="CHO-K1"))
        state.add_entity(_entity("cell_cho_k1_oatp1c1", "CellLineSample", name="CHO-K1 OATP1C1"))
        engine = _engine(state)

        pipeline_mod._draft_entities(engine)

        assert len(order) == 2, f"expected one call per cell line, got {len(order)}"
        assert order[0] != order[1], (
            "both cell lines were sent a byte-identical prompt — the model cannot "
            "tell which one it is describing"
        )

        # Each prompt must name the entity it is about, and must NOT be dominated
        # by the sibling's identity.
        contexts = dict(zip(["cell_cho_k1", "cell_cho_k1_oatp1c1"], order, strict=True))
        assert "CHO-K1 OATP1C1" in contexts["cell_cho_k1_oatp1c1"]
        assert "cell_cho_k1_oatp1c1" not in contexts["cell_cho_k1"], (
            "the parental line's prompt names its transfected sibling — this is "
            "exactly how the transfection description leaked onto the parental line"
        )

    def test_a_field_heavy_entity_cannot_balloon_the_prompt(self) -> None:
        """The identity block is bounded in total, not just per field.

        This fix exists partly to stop the drafter burning tokens; an entity
        carrying many long fields must not spend them back. Per-field truncation
        alone does not bound the block — twenty fields at the per-field cap is
        still a large prompt.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        entity = _entity(
            "as1",
            "Assay",
            **{f"field_{i}": ("x" * 400) for i in range(40)},
        )
        context = pipeline_mod._entity_draft_context(entity, "SHARED DIGEST")

        identity = context.split("SHARED DIGEST", 1)[1]
        assert len(identity) <= pipeline_mod._ENTITY_CONTEXT_MAX_CHARS + 200, (
            f"the per-entity block is unbounded ({len(identity)} chars) — a "
            "field-heavy entity would spend back what the skip guard saves"
        )
        # Identity survives the truncation: it is emitted before the field dump.
        assert "as1" in context
        assert "Assay" in context

    def test_the_shared_crate_context_is_still_present(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Per-entity identity is ADDED to the crate digest, not swapped for it.

        The leaf still needs the surrounding document to have anything to say;
        the fix must not starve it down to just an id.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        seen: list[str] = []

        def fake_leaf(entity_type, context, *, overrides=None, usage_sink=None):
            seen.append(context)
            return {}

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", fake_leaf)

        state = CrateState()
        state.metadata.title = "TPO inhibition dose-response screen"
        state.metadata.description = "A cell-based in vitro TPO inhibition assay."
        state.add_entity(_entity("as1", "Assay", name="As"))
        engine = _engine(state)

        pipeline_mod._draft_entities(engine)

        assert seen, "the leaf must still be called for a describable entity"
        assert "TPO inhibition" in seen[0], (
            "the shared crate context was dropped; the leaf has nothing to draft from"
        )


class TestGatherContextMetadataFirst:
    """`_gather_context` must feed the structured METADATA files to the leaf, not
    only filenames + the paper (Issue #179 / real S-VHPS26 run).

    Root cause (pre-fix): the 8000-char body budget is ONE shared pool and the
    per-file cap is ``min(remaining, _MAX_CONTEXT_CHARS)``, so the FIRST file read
    can consume the whole budget. Files are processed in plain scan order with no
    metadata prioritization, so a large early bulk-data file zeroes the budget and
    the richest structured metadata — a BioStudies ``<acc>.json`` export, an
    assay-metadata ``.xlsx``, a SOP ``.docx`` — never reaches ``extract_plan``.

    The fix: (a) order metadata-bearing files FIRST, and (b) cap each file fairly
    so no single early file starves the rest.
    """

    _JSON_SENTINEL = "BIOSTUDIES_ACCESSION_SENTINEL_S_VHPS26"
    _META_XLSX_SENTINEL = "ASSAY_METADATA_SHEET_SENTINEL"
    _DATA_SENTINEL = "BULK_DATA_FILE_BODY_SENTINEL"

    def _state_with_files(self, root: Path) -> CrateState:
        """A scanned-file inventory that reproduces the S-VHPS26 ordering trap.

        A LARGE bulk-data ``.xlsx`` is listed FIRST (its stubbed body alone exceeds
        the whole 8000-char budget), followed by the metadata-bearing files whose
        bodies carry unique sentinels. All paths are inside *root* (an approved
        scan root) so the fail-closed containment guard admits them.
        """
        from builder.agents.pipeline.pipeline import _MAX_CONTEXT_CHARS

        state = CrateState()
        state.metadata.title = "S-VHPS26 metadata-first context"
        state.approved_scan_roots = {str(root)}

        names = [
            "aaa_big_data.xlsx",  # bulk data, sorts FIRST alphabetically
            "S-VHPS26.json",  # BioStudies accession export (metadata)
            "assay_metadata.xlsx",  # assay-metadata sheet (metadata)
            "protocol_SOP.docx",  # SOP doc
            "zzz_more_data.xlsx",  # more bulk data, sorts LAST
        ]
        # Bodies: the FIRST file's body alone overflows the total budget; the
        # metadata files carry unique sentinel tokens near their start.
        self._bodies = {
            "aaa_big_data.xlsx": (self._DATA_SENTINEL + " ") * (_MAX_CONTEXT_CHARS * 2),
            "S-VHPS26.json": self._JSON_SENTINEL + ' {"accession": "S-VHPS26"}',
            "assay_metadata.xlsx": self._META_XLSX_SENTINEL + " | cell line | dose",
            "protocol_SOP.docx": "Standard operating procedure heading.",
            "zzz_more_data.xlsx": (self._DATA_SENTINEL + " ") * (_MAX_CONTEXT_CHARS * 2),
        }
        for name in names:
            p = root / name
            p.write_text("placeholder")  # real file so _contain resolves it
            state.scanned_files.append(
                FileClassification(
                    path=str(p),
                    filename=name,
                    size=len(self._bodies[name]),
                    mime_type="application/octet-stream",
                    first_rows=None,  # forces the body-read path for every file
                )
            )
        return state

    def _stub_reader(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Stub the body reader to return the sized bodies keyed by filename."""
        import builder.tools.file_readers as fr_mod

        def fake_read_file(path, *args, **kwargs):
            return self._bodies.get(Path(path).name)

        monkeypatch.setattr(fr_mod, "read_file", fake_read_file)

    def test_biostudies_json_sentinel_reaches_context(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """The BioStudies ``<acc>.json`` body MUST reach the gathered context even
        though a much larger bulk-data file is scanned first.

        Pre-fix this FAILS: the first file (``aaa_big_data.xlsx``) exhausts the
        whole 8000-char budget, so the JSON sentinel never appears.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._stub_reader(monkeypatch)
        engine = _engine(self._state_with_files(tmp_path))

        context = pipeline_mod._gather_context(engine)

        assert self._JSON_SENTINEL in context, (
            "the BioStudies .json metadata body must reach the leaf context"
        )

    def test_metadata_body_precedes_bulk_data_body(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """A metadata-named file's body must appear BEFORE a bulk data file's body."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._stub_reader(monkeypatch)
        engine = _engine(self._state_with_files(tmp_path))

        context = pipeline_mod._gather_context(engine)

        # The metadata .xlsx sentinel must appear, and before any bulk-data body.
        meta_idx = context.find(self._META_XLSX_SENTINEL)
        data_idx = context.find(self._DATA_SENTINEL)
        assert meta_idx != -1, "the assay-metadata file body must reach the context"
        if data_idx != -1:
            assert meta_idx < data_idx, (
                "metadata-named file body must precede a bulk data file body"
            )

    def test_total_budget_ceiling_is_respected(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """The overall 8000-char body budget stays the ceiling — no single file
        (and not the fair per-file slices summed) blows it."""
        import builder.agents.pipeline.pipeline as pipeline_mod
        from builder.agents.pipeline.pipeline import _MAX_CONTEXT_CHARS

        self._stub_reader(monkeypatch)
        engine = _engine(self._state_with_files(tmp_path))

        context = pipeline_mod._gather_context(engine)

        # Total body content (sentinels are a proxy: the two huge data files alone
        # would each exceed the budget if uncapped) stays bounded. Allow headroom
        # for the non-body scaffolding (title line, filename headers).
        assert len(context) <= _MAX_CONTEXT_CHARS + 2000


class TestMaterializePlan:
    """Stage B (`_materialize_plan`) turns the extracted candidate plan into real
    domain entities via the deterministic composites (Issue #179 task 2b-B).

    `extract_plan` (the whole-document leaf) is STUBBED at its pipeline import
    site — no model, no network — and the composites' own lookups
    (`lookup_compound` / `verify_identifier` / `lookup_aop`) are stubbed too, so
    these tests are fully offline. The contract under test:

    * a canned plan materializes the expected MolecularEntity / CellLineSample /
      LabProcess chain / AOP subgraph / Person / Publication entities in state;
    * with NO provider configured, the step is a STRICT no-op;
    * with a provider but NO usable context, the step is a strict no-op;
    * D5 — no identifier from the plan is ever set on an entity; identifiers come
      only from the composites' lookups/verification;
    * idempotent — running the spine twice produces no duplicates;
    * `run_pipeline` still reaches `{base, isa, tox}` conformance with a plan.
    """

    _PLAN: dict = {
        "study": {"name": "TPO inhibition study", "description": "A TPO assay."},
        "compounds": [
            {"name": "Methimazole", "role": "test"},
            {"name": "Sodium iodide", "role": "control"},
        ],
        "cell_lines": [{"name": "FRTL-5"}],
        "protocols": [
            {
                "name": "Amplex Red TPO activity readout",
                "description": "Fluorometric TPO activity assay protocol.",
                "process_hint": "EndpointReadout",
            }
        ],
        "process_chain": [
            # Exposure / EndpointReadout / DataAnalysis each MUST carry at least
            # one schema:additionalProperty under the tox profile, and `_pv` no
            # longer publishes "unknown" to satisfy it — so a plan whose crate is
            # expected to CONFORM has to supply a real parameter per step.
            {"process_type": "CellCulture", "name": "Seed cells"},
            {
                "process_type": "Exposure",
                "name": "Dose",
                "parameters": {"duration": "24 hours"},
            },
            {
                "process_type": "EndpointReadout",
                "name": "Read TPO",
                "parameters": {"detection_instrument": "Plate reader"},
            },
            {
                "process_type": "DataAnalysis",
                "name": "Fit dose-response",
                "parameters": {"data_processing": "Four-parameter logistic fit"},
            },
        ],
        "aops": [{"aop_id": "610"}],
        "people": [{"name": "Ada Lovelace", "affiliation_name": "Analytical Engine"}],
        "publications": [{"title": "On TPO inhibition in vitro"}],
        "files": [{"path": "data/raw.csv", "role": "raw"}],
        "notes": "Confirm the control compound role.",
    }

    def _titled_state(self) -> CrateState:
        state = CrateState()
        state.metadata.title = "TPO inhibition dose-response screen"
        state.metadata.description = "A cell-based in vitro TPO inhibition assay."
        return state

    def _enable_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")

    def _stub_extract_plan(
        self, monkeypatch: pytest.MonkeyPatch, plan: dict | None = None
    ) -> list[str]:
        """Patch the pipeline's `extract_plan` shim to return a canned plan."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        seen: list[str] = []

        def fake_extract_plan(context, *, overrides=None, usage_sink=None):
            seen.append(context)
            return dict(self._PLAN if plan is None else plan)

        monkeypatch.setattr(pipeline_mod, "extract_plan", fake_extract_plan)
        return seen

    def _stub_lookups(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Stub every network lookup the composites would otherwise hit."""
        import builder.tools.composites as composites_mod
        from builder.tools import lookups as tool_lookups
        from builder.tools._resolve_cache import compound_cache

        # The compound resolution cache is process-global; a prior test can pre-cache
        # these names and short-circuit the lookup, masking this stub's per-name CID.
        # Clear it so the stub always runs fresh (xdist-safe).
        compound_cache.clear()

        # resolve_compound -> lookup_compound (imported into composites' namespace).
        # A DISTINCT CID per compound name: the dedup-by-chemical-identity path
        # (Issue #179) collapses two names that resolve to the SAME identity into one
        # MolecularEntity, so two DISTINCT plan compounds must carry distinct CIDs
        # to stay distinct nodes. CAS stays constant (`60-56-0`) so the D5 exact-value
        # assertion is preserved; only the (node-id-bearing) CID varies per name.
        _cids = {"Methimazole": "1349907", "Sodium iodide": "5238"}

        def fake_lookup_compound(name):
            return {
                "found": True,
                "data": {
                    "cas": "60-56-0",
                    "pubchem_cid": _cids.get(name, "999999"),
                    "smiles": "C1=CN(C(=S)N1)C",
                    "source": "pubchem",
                },
                "error": None,
            }

        # verify_identifier marks the field verified without touching the network.
        def fake_verify_identifier(state, entity_id, field):
            ent = state.get_entity(entity_id)
            if ent is not None:
                ent.set_field_status(field, "verified", "lookup")
            return {"verified": True, "entity_id": entity_id, "field": field, "message": "ok"}

        # materialize_aop_subgraph -> lookup_aop (imported lazily from tool_lookups).
        def fake_lookup_aop(aop_id):
            iri = f"https://aopwiki.org/aops/{aop_id}"
            mie = "https://aopwiki.org/events/1"
            ao = "https://aopwiki.org/events/2"
            ker = "https://aopwiki.org/relationships/1"
            return {
                "found": True,
                "data": {
                    "aop": {
                        "@id": iri,
                        "@type": "AdverseOutcomePathway",
                        "name": f"AOP {aop_id}",
                        "identifier": str(aop_id),
                        "url": iri,
                        "has_molecular_initiating_event": [{"@id": mie}],
                        "has_adverse_outcome": [{"@id": ao}],
                        "has_key_event_relationship": [{"@id": ker}],
                    },
                    "events": [
                        {
                            "@id": mie,
                            "@type": "KeyEvent",
                            # A REAL event name (#382): the Assay -> Key Event link
                            # matches the plan's `measured_event_name` against these
                            # names, so a placeholder "MIE" could not exercise it.
                            "name": "Mitochondrial dysfunction",
                            "eventType": "Molecular Initiating Event",
                        },
                        {
                            "@id": ao,
                            "@type": "KeyEvent",
                            "name": "AO",
                            "eventType": "Adverse Outcome",
                        },
                    ],
                    "relationships": [
                        {
                            "@id": ker,
                            "@type": "KeyEventRelationship",
                            "upstream_event": {"@id": mie},
                            "downstream_event": {"@id": ao},
                        },
                    ],
                },
                "error": None,
            }

        # resolve_publication -> search_works_by_title. Default: NO candidates, so
        # the default-plan publication stays deferred (D5). Tests that need a
        # confident match override this stub.
        def fake_search_works_by_title(title):
            return []

        monkeypatch.setattr(composites_mod, "lookup_compound", fake_lookup_compound)
        monkeypatch.setattr(composites_mod, "verify_identifier", fake_verify_identifier)
        monkeypatch.setattr(composites_mod, "search_works_by_title", fake_search_works_by_title)
        monkeypatch.setattr(tool_lookups, "lookup_aop", fake_lookup_aop)

    def _by_type(self, engine: AgentEngine, type_name: str) -> list[Entity]:
        return [e for e in engine.state.list_entities() if e.type == type_name]

    def test_materializes_plan_entities(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        seen = self._stub_extract_plan(monkeypatch)
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        # Scaffold first so the assay/study ids exist for the chain/aop wiring.
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        # The leaf saw the real gathered context (the crate title).
        assert seen and any("TPO inhibition" in ctx for ctx in seen)

        # Compounds → MolecularEntity (one per plan compound).
        chems = self._by_type(engine, "MolecularEntity")
        assert {c.fields.get("name") for c in chems} == {"Methimazole", "Sodium iodide"}

        # Cell line → CellLineSample.
        cells = self._by_type(engine, "CellLineSample")
        assert [c.fields.get("name") for c in cells] == ["FRTL-5"]

        # Process chain → 4 LabProcess steps wired to the assay.
        procs = self._by_type(engine, "LabProcess")
        ptypes = {p.fields.get("process_type") for p in procs}
        assert ptypes == {"CellCulture", "Exposure", "EndpointReadout", "DataAnalysis"}

        # Protocol → LabProtocol minted from the name/description (D5: no id), and
        # linked to the EndpointReadout process it governs (executesLabProtocol).
        protos = self._by_type(engine, "LabProtocol")
        assert [p.fields.get("name") for p in protos] == ["Amplex Red TPO activity readout"]
        proto_id = protos[0].entity_id
        readout = next(p for p in procs if p.fields.get("process_type") == "EndpointReadout")
        labprotocol_ref = readout.fields.get("labprotocol")
        ref_id = (
            labprotocol_ref.get("@id") if isinstance(labprotocol_ref, dict) else labprotocol_ref
        )
        assert str(ref_id).lstrip("#") == proto_id
        assert result["protocols"] >= 1

        # AOP → AdverseOutcomePathway subgraph.
        assert self._by_type(engine, "AdverseOutcomePathway")
        assert self._by_type(engine, "KeyEvent")

        # Person from the name, with a deterministic given/family split so it is
        # ISA-conformant (a non-empty given name is REQUIRED). No ORCID (D5).
        people = self._by_type(engine, "Person")
        ada = next(p for p in people if p.fields.get("name") == "Ada Lovelace")
        assert ada.fields.get("givenName") == "Ada"
        assert ada.fields.get("familyName") == "Lovelace"
        assert not ada.fields.get("orcid")

        # Publications are DEFERRED (title-only cannot be ISA-conformant without a
        # fabricated DOI — D5), so no Publication entity is minted; the title is
        # surfaced for a later DOI lookup instead.
        assert self._by_type(engine, "Publication") == []
        assert "On TPO inhibition in vitro" in result["publications_deferred"]
        assert result["publications"] == 0

        # The result is informative (per-section counts).
        assert result["compounds"] >= 2
        assert result["cell_lines"] >= 1
        assert result["processes"] >= 1
        assert result["aops"] >= 1
        assert result["people"] >= 1

    def test_protocol_minted_and_linked_to_process(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """#222: a plan protocol mints a LabProtocol linked to a process (D5).

        A plan carries a protocol NAME/description only (no identifier). The spine
        mints exactly one LabProtocol and wires it onto the LabProcess it governs
        via the ``labprotocol`` ref (``executesLabProtocol`` at build time).
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        plan = {
            "protocols": [{"name": "MTT viability protocol", "description": "MTT assay."}],
            "process_chain": [
                {"process_type": "Exposure", "name": "Dose"},
                {"process_type": "EndpointReadout", "name": "Read viability"},
            ],
        }
        self._stub_extract_plan(monkeypatch, plan)
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        protos = self._by_type(engine, "LabProtocol")
        assert len(protos) == 1
        assert protos[0].fields.get("name") == "MTT viability protocol"
        # D5: a protocol minted from a name carries no fabricated identifier.
        assert not protos[0].fields.get("identifier")
        assert result["protocols"] == 1

        # The protocol is linked to at least one LabProcess.
        proc_refs = [
            p.fields.get("labprotocol")
            for p in self._by_type(engine, "LabProcess")
            if p.fields.get("labprotocol")
        ]
        assert proc_refs, "the protocol must be linked to a process"

    def test_confident_publication_match_mints_entity(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """#224: a confident title match mints a ScholarlyArticle, not deferred."""
        import builder.agents.pipeline.pipeline as pipeline_mod
        import builder.tools.composites as composites_mod

        self._enable_provider(monkeypatch)
        title = "On TPO inhibition in vitro"
        self._stub_extract_plan(monkeypatch, {"publications": [{"title": title}]})
        self._stub_lookups(monkeypatch)

        # Confident Crossref candidate (exact title, high score) → commit the DOI.
        def fake_search_works_by_title(query):
            return [{"doi": "10.1234/tpo", "title": title, "score": 99.0}]

        # The drafter re-looks the DOI up; canned article data (no network).
        def fake_lookup_doi(doi):
            return {
                "found": True,
                "data": {
                    "identifier": "10.1234/tpo",
                    "name": title,
                    "author": [{"givenName": "Ada", "familyName": "Lovelace"}],
                },
                "error": None,
            }

        monkeypatch.setattr(composites_mod, "search_works_by_title", fake_search_works_by_title)
        monkeypatch.setattr(composites_mod, "lookup_doi", fake_lookup_doi)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        pubs = self._by_type(engine, "Publication")
        assert len(pubs) == 1
        assert pubs[0].fields.get("identifier") == "10.1234/tpo"
        assert result["publications"] >= 1
        assert title not in result["publications_deferred"]

    def test_no_confident_publication_match_stays_deferred(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """#224: no confident match → no entity, title stays deferred (D5)."""
        import builder.agents.pipeline.pipeline as pipeline_mod
        import builder.tools.composites as composites_mod

        self._enable_provider(monkeypatch)
        title = "An unfindable paper"
        self._stub_extract_plan(monkeypatch, {"publications": [{"title": title}]})
        self._stub_lookups(monkeypatch)

        # No candidate clears the confidence gate → resolve_publication returns
        # ok=False and mints nothing; the spine must NOT fabricate a DOI.
        def fake_search_works_by_title(query):
            return [{"doi": "10.9/unrelated", "title": "A different paper", "score": 99.0}]

        def boom_lookup_doi(doi):  # pragma: no cover - must not be reached
            raise AssertionError("lookup_doi must not run without a confident match")

        monkeypatch.setattr(composites_mod, "search_works_by_title", fake_search_works_by_title)
        monkeypatch.setattr(composites_mod, "lookup_doi", boom_lookup_doi)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        assert self._by_type(engine, "Publication") == []
        assert result["publications"] == 0
        assert title in result["publications_deferred"]

    def test_d5_no_fabricated_identifiers_from_plan(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """D5: identifiers come from the composites' lookups, never the plan.

        The plan carries only names. A MolecularEntity's `cas`/`pubchem_cid` must
        be the LOOKED-UP value (here the stubbed `60-56-0` / `1349907`), and a
        fabricated DOI an adversarial plan smuggles in must never land on any
        entity (the title-only publication is deferred, not materialized).
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        # A plan that adversarially tries to smuggle identifiers (should be ignored
        # by the materialization — names only are passed to the composites).
        plan = {
            "compounds": [{"name": "Methimazole", "role": "test", "cas": "FAKE-CAS"}],
            "publications": [{"title": "Some paper", "doi": "10.0/FAKE"}],
        }
        self._stub_extract_plan(monkeypatch, plan)
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        chem = self._by_type(engine, "MolecularEntity")[0]
        # The CAS is the looked-up value, NOT the plan's fabricated one.
        assert chem.fields.get("cas") == "60-56-0"
        assert chem.fields.get("cas") != "FAKE-CAS"

        # The publication is deferred (title surfaced) and no entity carries the
        # plan's fabricated DOI anywhere in state.
        assert "Some paper" in result["publications_deferred"]
        for ent in engine.state.list_entities():
            assert ent.fields.get("identifier") != "10.0/FAKE"
            assert ent.fields.get("doi") != "10.0/FAKE"

    # --- the Assay -> AOP Key Event link (#382) ------------------------------
    #
    # Materializing a pathway used to be the end of the AOP section: the crate
    # listed every KeyEvent of AOP 610 and never said which one the assay
    # measured, so the assay-to-mechanism edge was absent from every crate the
    # tool had ever produced. The plan now carries the event NAME and the spine
    # feeds it to `link_assay_to_key_event`, which commits the id AOP-Wiki gave
    # it. One claim, then controls: the link must be CAUSED by
    # `measured_event_name` flowing through the tool — so it must NOT appear when
    # the plan is silent or names an event this pathway does not have, and no id
    # the plan carries may ever reach the crate.

    def _assay_key_event(self, engine: AgentEngine) -> Any:
        """The Assay's stored `keyEvent` reference id, or None."""
        assay = self._by_type(engine, "Assay")[0]
        ref = assay.fields.get("keyEvent")
        return ref.get("@id") if isinstance(ref, dict) else ref

    def test_plan_measured_event_wires_assay_to_key_event(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(
            monkeypatch,
            {
                "aops": [
                    {"aop_id": "610", "measured_event_name": "Mitochondrial dysfunction"}
                ]
            },
        )
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        # The committed reference is the IRI the LOOKUP STUB produced for the
        # event of that name — the test never writes it onto an entity, and the
        # plan carries only the name.
        assert self._assay_key_event(engine) == "https://aopwiki.org/events/1"
        assert result["key_events"] == 1
        # The MIT slot is `Assay:keyEvent` and MIT scoring keys on the raw state
        # field name, so the snake_case spelling would score as unfilled.
        assert "key_event" not in self._by_type(engine, "Assay")[0].fields

    def test_unnamed_event_leaves_assay_unlinked(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Honesty control: the default plan names no event, so nothing is linked.

        This is today's behaviour and must stay: it proves the test above is
        driven by `measured_event_name` rather than by anything the AOP section
        does unconditionally.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch)  # the default _PLAN: aop_id only
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        # The subgraph IS fully materialized; only the assay edge is missing.
        assert self._by_type(engine, "AdverseOutcomePathway")
        assert self._by_type(engine, "KeyEvent")
        assert "keyEvent" not in self._by_type(engine, "Assay")[0].fields
        assert result["key_events"] == 0

    def test_plan_event_name_matching_nothing_leaves_assay_unlinked(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Honesty control: a name AOP-610 does not carry links nothing, quietly.

        "TPO inhibition" vs "Thyroperoxidase, Inhibition" is the real abbreviation
        gap this refusal exists for — no string matcher may bridge it, and being
        unable to is a normal outcome, not a spine failure.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(
            monkeypatch,
            {
                "aops": [
                    {"aop_id": "610", "measured_event_name": "Thyroperoxidase, Inhibition"}
                ]
            },
        )
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)  # must not raise

        assert "keyEvent" not in self._by_type(engine, "Assay")[0].fields
        assert result["key_events"] == 0
        assert result["aops"] == 1  # the pathway itself still landed

    def test_plan_cannot_supply_an_event_identifier(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """D5: an event id an adversarial model attaches never reaches the crate.

        The plan is passed through the REAL `_strip_plan_identifiers` — the same
        scrub `extract_plan` applies to the model's output — before the spine sees
        it, so this exercises the guard rather than asserting the stub's own shape.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod
        from builder.agents.pipeline.leaves import _strip_plan_identifiers

        self._enable_provider(monkeypatch)
        raw_plan = {
            "aops": [
                {
                    "aop_id": "610",
                    "event_id": "999",
                    "measured_event_name": "Mitochondrial dysfunction",
                }
            ]
        }
        self._stub_extract_plan(monkeypatch, _strip_plan_identifiers(raw_plan))
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)

        # The NAME still resolves through the lookup, so the link is made — from
        # the lookup's id, never the plan's.
        assert self._assay_key_event(engine) == "https://aopwiki.org/events/1"
        for ent in engine.state.list_entities():
            assert "999" not in ent.entity_id
            for value in ent.fields.values():
                assert "999" not in str(value)

    def test_the_scrub_removes_event_ids_and_keeps_the_event_name(self) -> None:
        """The scrub itself, not just its effect downstream (D5).

        ``_materialize_plan`` reads only ``aop_id`` and ``measured_event_name``, so
        an unstripped ``event_id`` cannot reach the crate through it — which means
        the end-to-end D5 test above stays green even with ``event_id`` / ``ke_id``
        / ``mie_id`` removed from :data:`_PLAN_IDENTIFIER_FIELDS`, and pins nothing
        about the guard. Asserting the stripper directly is what makes those three
        entries load-bearing: they are defense in depth against a FUTURE reader of
        ``aops[]``, and defense nothing exercises rots.
        """
        from builder.agents.pipeline.leaves import _strip_plan_identifiers

        scrubbed = _strip_plan_identifiers(
            {
                "aops": [
                    {
                        "aop_id": "610",
                        "event_id": "999",
                        "ke_id": "888",
                        "mie_id": "777",
                        "measured_event_name": "Mitochondrial dysfunction",
                    }
                ]
            }
        )

        item = scrubbed["aops"][0]
        assert set(item) == {"aop_id", "measured_event_name"}, (
            "an AOP-Wiki event id must not survive the scrub; the name must"
        )
        assert item["measured_event_name"] == "Mitochondrial dysfunction"

    def test_plan_schema_offers_no_event_identifier_field(self) -> None:
        """The new slot must never grow an id sibling the model could fill (D5)."""
        from builder.agents.pipeline.leaves import (
            _PLAN_IDENTIFIER_FIELDS,
            _plan_schema,
        )

        names: set[str] = set()

        def walk(node: Any) -> None:
            if isinstance(node, dict):
                properties = node.get("properties")
                if isinstance(properties, dict):
                    names.update(properties)
                for value in node.values():
                    walk(value)
            elif isinstance(node, list):
                for item in node:
                    walk(item)

        walk(_plan_schema())
        # The slot that must exist, spelled as a name...
        assert "measured_event_name" in names
        # ...and no identifier-shaped property anywhere in the schema.
        assert not (names & _PLAN_IDENTIFIER_FIELDS)

    def test_no_provider_skips_only_the_plan_driven_sections(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """With no provider the PLAN-DRIVEN sections are a strict no-op (extract_plan
        is never called), but the deterministic chain + file steps still run (#262).
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: None)

        def boom(*args, **kwargs):  # pragma: no cover - must not run
            raise AssertionError("extract_plan must not run without a provider")

        monkeypatch.setattr(pipeline_mod, "extract_plan", boom)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        # The plan-driven sections mint nothing (no model, no lookups).
        assert result.get("compounds") == 0
        assert result.get("cell_lines") == 0
        assert result.get("aops") == 0
        assert result.get("people") == 0
        assert result.get("publications") == 0
        assert self._by_type(engine, "MolecularEntity") == []
        assert self._by_type(engine, "AdverseOutcomePathway") == []
        # But the deterministic process chain DID run (the #262 contract).
        assert result.get("processes") == 4

    def test_no_context_skips_only_the_plan_driven_sections(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """With a provider but no usable context, extract_plan is never called and
        the plan-driven sections mint nothing; the deterministic chain still runs.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)

        def boom(*args, **kwargs):  # pragma: no cover - must not run
            raise AssertionError("extract_plan must not run without context")

        monkeypatch.setattr(pipeline_mod, "extract_plan", boom)

        # Untitled / undescribed / unscanned crate carries no usable context.
        engine = _engine(CrateState())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        assert result.get("compounds") == 0
        assert result.get("aops") == 0
        assert self._by_type(engine, "MolecularEntity") == []
        # No scanned files ⇒ nothing attached, but the chain still ran.
        assert result.get("files") == 0
        assert result.get("processes") == 4

    def test_idempotent_no_duplicates(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch)
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)
        counts_1 = {
            t: len(self._by_type(engine, t))
            for t in (
                "MolecularEntity",
                "CellLineSample",
                "LabProcess",
                "AdverseOutcomePathway",
                "KeyEvent",
                "Person",
                "Publication",
            )
        }
        # Re-run on the SAME engine — composites are idempotent, so no dups.
        pipeline_mod._materialize_plan(engine)
        counts_2 = {
            t: len(self._by_type(engine, t))
            for t in (
                "MolecularEntity",
                "CellLineSample",
                "LabProcess",
                "AdverseOutcomePathway",
                "KeyEvent",
                "Person",
                "Publication",
            )
        }
        assert counts_1 == counts_2

    @staticmethod
    def _affiliation_ref_id(person: Entity) -> str | None:
        """The bare entity_id a Person's ``affiliation`` field points at.

        ``set_fields`` stores the reference verbatim (a bare id or an
        ``{"@id": …}`` object), so normalize both to the leading-``#``-stripped id.
        """
        value = person.fields.get("affiliation")
        ref = value.get("@id") if isinstance(value, dict) else value
        return str(ref).lstrip("#") if ref else None

    def test_person_affiliation_mints_organization_and_links(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Lane 1 — a plan ``people[].affiliation_name`` mints an Organization and
        wires the Person's ``affiliation`` onto it.

        ``extract_plan`` already surfaces ``affiliation_name`` (leaves.py), but the
        deterministic materialize people-loop previously dropped it: the crate ended
        up with ZERO Organization entities and the Person carried no affiliation.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch)
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)

        # (a) an Organization named after the plan affiliation exists in state.
        orgs = self._by_type(engine, "Organization")
        analytical = [o for o in orgs if o.fields.get("name") == "Analytical Engine"]
        assert len(analytical) == 1, "exactly one Organization should be minted"

        # (b) the Person's affiliation references that Organization's id.
        ada = next(
            p for p in self._by_type(engine, "Person") if p.fields.get("name") == "Ada Lovelace"
        )
        assert self._affiliation_ref_id(ada) == analytical[0].entity_id

    def test_shared_affiliation_is_deduplicated(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Two people with the SAME affiliation_name produce ONE Organization,
        both referencing it (no duplicate orgs)."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        plan = dict(self._PLAN)
        plan["people"] = [
            {"name": "Ada Lovelace", "affiliation_name": "Analytical Engine"},
            {"name": "Charles Babbage", "affiliation_name": "Analytical Engine"},
        ]

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch, plan)
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)

        orgs = [
            o
            for o in self._by_type(engine, "Organization")
            if o.fields.get("name") == "Analytical Engine"
        ]
        assert len(orgs) == 1, "shared affiliation must mint exactly one Organization"

        org_id = orgs[0].entity_id
        people = {p.fields.get("name"): p for p in self._by_type(engine, "Person")}
        assert self._affiliation_ref_id(people["Ada Lovelace"]) == org_id
        assert self._affiliation_ref_id(people["Charles Babbage"]) == org_id

    def test_person_without_affiliation_mints_no_organization(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """A person with no affiliation_name mints no Organization and gets no
        affiliation field — no regression, no fabrication (D5)."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        plan = dict(self._PLAN)
        plan["people"] = [{"name": "Ada Lovelace"}]

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch, plan)
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)

        assert self._by_type(engine, "Organization") == []
        ada = next(
            p for p in self._by_type(engine, "Person") if p.fields.get("name") == "Ada Lovelace"
        )
        assert not ada.fields.get("affiliation")

    def test_run_pipeline_with_plan_reaches_conformance(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        from builder.agents.pipeline.pipeline import run_pipeline

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch)
        self._stub_lookups(monkeypatch)
        # Keep the field-enrichment leaf a no-op (it is separately tested) so this
        # test isolates materialization + the existing build/fix path.
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", lambda *a, **k: {})

        engine = _engine(self._titled_state())
        result = run_pipeline(engine)

        assert result["conformance"] == {"base": True, "isa": True, "tox": True}
        assert result["ok"] is True
        # The materialized plan is reflected in the result trace.
        assert "materialized" in result
        assert result["materialized"]["compounds"] >= 2


def _types(node: dict) -> set[str]:
    """The @type(s) of a graph node as a set of strings."""
    t = node.get("@type")
    if isinstance(t, list):
        return {str(x) for x in t}
    return {str(t)} if t is not None else set()


class TestMaterializeLinksResolvedEntities(TestMaterializePlan):
    """Issue #273 — resolved compounds and the cell line must be WIRED into the
    provenance, not left as orphans.

    The materialize path already RESOLVES the right MolecularEntity / CellLineSample
    entities, but before #273 it never linked them into the process chain, so the
    exported crate flagged them all as ``⚠ orphan``. The fix wires them with the
    canonical ISA-Tox reference fields:

    * each resolved ``MolecularEntity`` → the Exposure LabProcess via the
      ``chemicals`` ref field. ISA forbids a MolecularEntity as a process object
      (objects MUST be File/Sample/BioSample), so the compound is connected THROUGH
      the Exposure's CSVW condition table (``schema:about`` → MolecularEntity) and,
      at a glance, on the Study via ``schema:mentions`` (the ``chemicals`` Study
      mention).
    * the resolved ``CellLineSample`` → the CellCulture LabProcess via the
      ``cell_line`` ref field (its consumed input), replacing the synthesized
      generic ``..._input`` placeholder; also surfaced on the Study via
      ``cell_lines`` (``biologicalModels``, an alias of ``schema:mentions``).

    Reachability is asserted against the *built* ``@graph`` (the exact assembly
    ``build_and_validate`` uses): every resolved MolecularEntity / CellLineSample
    node must be referenced by at least one other node, i.e. orphan count → 0.
    """

    # NOTE: the base ``_stub_lookups`` already hands a DISTINCT CID per compound
    # name (and clears the process-global compound cache), so the two plan
    # compounds resolve to two distinct MolecularEntity nodes — exactly what this
    # #273 reachability test needs. No override required.

    @staticmethod
    def _referenced_ids(graph: list[dict]) -> set[str]:
        """Every @id referenced by ANY node in *graph*.

        Walks each node's property values for ``{"@id": ...}`` references (scalar,
        list, or nested) EXCEPT the node's own ``@id``/``@type``, so a node that
        only "references" itself is not counted as referenced.
        """
        referenced: set[str] = set()

        def _collect(value: object) -> None:
            if isinstance(value, dict):
                ref = value.get("@id")
                if isinstance(ref, str):
                    referenced.add(ref)
                else:
                    for v in value.values():
                        _collect(v)
            elif isinstance(value, list):
                for item in value:
                    _collect(item)

        for node in graph:
            for key, value in node.items():
                if key in ("@id", "@type"):
                    continue
                _collect(value)
        return referenced

    def _built_graph(self, engine: AgentEngine) -> list[dict]:
        """The built crate's ``@graph`` — the same assembly build_and_validate uses."""
        from builder.tools.builder import assemble_crate

        crate = assemble_crate(
            engine.state,
            output_dir=None,
            materialize_payload=False,
            include_all_scanned=False,
        )
        return crate.metadata.generate()["@graph"]

    def test_compounds_and_cell_line_are_wired_not_orphaned(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch)
        self._stub_lookups(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)

        # The Exposure carries the resolved MolecularEntity ids in `chemicals`.
        procs = self._by_type(engine, "LabProcess")
        exposure = next(p for p in procs if p.fields.get("process_type") == "Exposure")
        chem_ids = {c.entity_id for c in self._by_type(engine, "MolecularEntity")}
        assert chem_ids, "no MolecularEntity resolved — test setup is wrong"
        wired_chems = exposure.fields.get("chemicals")
        wired_chem_ids = {
            (c.get("@id") if isinstance(c, dict) else c)
            for c in (wired_chems if isinstance(wired_chems, list) else [wired_chems])
        }
        wired_chem_ids = {str(c).lstrip("#") for c in wired_chem_ids if c}
        assert wired_chem_ids == {c.lstrip("#") for c in chem_ids}

        # The CellCulture references the actual cell-line Sample via `cell_line`
        # (not a synthesized generic `_input`).
        cell_culture = next(p for p in procs if p.fields.get("process_type") == "CellCulture")
        cell_ids = {c.entity_id for c in self._by_type(engine, "CellLineSample")}
        assert cell_ids, "no CellLineSample resolved — test setup is wrong"
        wired_cell = cell_culture.fields.get("cell_line")
        wired_cell_id = wired_cell.get("@id") if isinstance(wired_cell, dict) else wired_cell
        assert str(wired_cell_id).lstrip("#") in {c.lstrip("#") for c in cell_ids}

        # No resolved MolecularEntity / CellLineSample is an orphan in the BUILT
        # graph: every one is referenced by at least one other node (#273).
        graph = self._built_graph(engine)
        referenced = self._referenced_ids(graph)
        node_ids = {n.get("@id") for n in graph}

        # A CellLineSample builds as @type Sample discriminated by
        # additionalType="CellLine" (its intermediate derived samples are plain
        # Samples), so detect it via that discriminator, not the state class name.
        def _is_compound(n: dict) -> bool:
            return "MolecularEntity" in _types(n)

        def _is_cell_line(n: dict) -> bool:
            return n.get("additionalType") == "CellLine"

        domain_nodes = [n for n in graph if _is_compound(n) or _is_cell_line(n)]
        # Both MolecularEntity (the two plan compounds) AND the CellLineSample must
        # be present in the built graph (the resolver mints them; the compound node
        # @id is the verified identifier IRI, not the state entity_id).
        assert len([n for n in domain_nodes if _is_compound(n)]) == 2
        assert any(_is_cell_line(n) for n in domain_nodes)
        # …and NONE of them is an orphan: every one is referenced by another node.
        orphans = [n["@id"] for n in domain_nodes if n["@id"] not in referenced]
        assert orphans == [], f"orphaned resolved domain entities: {orphans}"
        # The Study surfaces every compound + the cell line via schema:mentions
        # (emitted under the context-aliased `chemicals` / `biologicalModels`
        # keys), so each is reachable from the backbone at a glance (#273).
        study_node = next(n for n in graph if str(n.get("@id", "")).startswith("#Study_"))

        def _ref_ids(prop: object) -> set[str]:
            items = prop if isinstance(prop, list) else [prop]
            ids = (m.get("@id") if isinstance(m, dict) else m for m in items)
            return {ref for ref in ids if isinstance(ref, str)}

        compound_node_ids = {n["@id"] for n in domain_nodes if _is_compound(n)}
        cell_node_ids = {n["@id"] for n in domain_nodes if _is_cell_line(n)}
        assert compound_node_ids <= _ref_ids(study_node.get("chemicals"))
        assert cell_node_ids <= _ref_ids(study_node.get("biologicalModels"))
        assert node_ids  # graph is non-empty

    def test_run_pipeline_with_links_preserves_conformance(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Wiring the compounds/cell line must NOT regress ISA + Tox conformance."""
        from builder.agents.pipeline.pipeline import run_pipeline

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch)
        self._stub_lookups(monkeypatch)
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", lambda *a, **k: {})

        engine = _engine(self._titled_state())
        result = run_pipeline(engine)

        assert result["conformance"] == {"base": True, "isa": True, "tox": True}
        assert result["ok"] is True


class TestMaterializeCompoundsFromFilenames:
    """#258 — the DEFAULT pipeline path must end-to-end turn compound-bearing
    DATA FILENAMES into MolecularEntities (the legacy ReAct path got 22; the
    default path got 0). This drives the REAL ``extract_plan`` leaf (its chat
    model faked offline to return the filename-derived compound NAMES the enriched
    prompt steers a model toward) through ``_materialize_plan`` with the compound
    lookups stubbed, so the whole filename → plan → ``resolve_compound`` →
    MolecularEntity chain is exercised with NO network.

    Contract:
    * count of MolecularEntities goes 0 → N for the representative input;
    * NAMES-ONLY (D5): only the plan NAME reaches ``resolve_compound``; the
      MolecularEntity's CAS/CID is the LOOKED-UP value, never invented from a
      filename or a plan field.
    """

    # The S-VHPS26 shape: compounds appear ONLY in the data filenames.
    _COMPOUND_FILES = (
        "S-VHPS26_P5_Silychristin+Verapamil.xlsx",
        "S-VHPS26_Diclofenac+BSP.xlsx",
    )
    _EXPECTED_NAMES = {"Silychristin", "Verapamil", "Diclofenac", "BSP"}

    def _enable_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")

    def _state_with_compound_files(self, tmp_path: Path) -> CrateState:
        """A titled state whose scanned data files name compounds in their stems."""
        state = CrateState()
        state.metadata.title = "S-VHPS26 transporter interaction screen"
        state.approved_scan_roots.add(str(tmp_path.resolve()))
        files: list[FileClassification] = []
        for name in self._COMPOUND_FILES:
            path = tmp_path / name
            path.write_bytes(b"PK\x03\x04 xlsx stub")  # body irrelevant — names in filenames
            files.append(
                FileClassification(
                    path=str(path),
                    filename=name,
                    size=path.stat().st_size,
                    mime_type=("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
                    # A tabular preview so _gather_context lists the filename
                    # WITHOUT a disk read (the filename carries the compound names).
                    first_rows=["well,value", "A1,1.0"],
                )
            )
        state.scanned_files = files
        return state

    def _fake_extract_plan_chat_model(self, monkeypatch: pytest.MonkeyPatch) -> list[str]:
        """Fake the REAL leaf's chat model so the genuine ``extract_plan`` code
        path runs (its enriched prompt is built and fed the gathered context),
        but the model deterministically returns the filename-derived compounds a
        correctly-steered model would. Records the prompt text the model saw so a
        test can assert the filenames actually reached the leaf."""
        import builder.agents.pipeline.leaves as leaves_mod

        seen_prompts: list[str] = []
        expected = self._EXPECTED_NAMES

        class _Runnable:
            def invoke(self, messages, *a, **k):
                for msg in messages:
                    content = getattr(msg, "content", "")
                    if isinstance(content, str):
                        seen_prompts.append(content)
                # A steered model proposes the compound NAMES it read off the
                # filenames — names only (D5: no identifiers).
                return {"compounds": [{"name": n, "role": "test"} for n in sorted(expected)]}

        class _Model:
            def with_structured_output(self, schema, *, include_raw=False, **k):
                return _Runnable()

        monkeypatch.setattr(leaves_mod, "_build_chat_model", lambda *a, **k: _Model())
        return seen_prompts

    # Per-name LOOKED-UP identifiers (offline) — the verified ids PubChem would
    # have returned for these names. Distinct per name so a test can assert the
    # right name resolved to the right (looked-up, not invented) CAS.
    _LOOKUP_BY_NAME: dict[str, dict[str, str]] = {
        "Silychristin": {"cas": "33889-69-9", "pubchem_cid": "441764"},
        "Verapamil": {"cas": "52-53-9", "pubchem_cid": "2520"},
        "Diclofenac": {"cas": "15307-86-5", "pubchem_cid": "3033"},
        "BSP": {"cas": "71-67-0", "pubchem_cid": "9568"},
    }

    def _stub_compound_lookup(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Stub ``resolve_compound``'s name->record lookup + verification offline,
        so a name resolves to a VERIFIED (looked-up) identifier with no network —
        the D5 path: the id comes from the lookup, never from the plan/filename."""
        import builder.tools.composites as composites_mod

        lookup_by_name = self._LOOKUP_BY_NAME

        def fake_lookup_compound(name):
            ids = lookup_by_name.get(str(name).strip())
            if ids is None:
                return {"found": False, "data": None, "error": "not found"}
            return {"found": True, "data": {**ids, "source": "pubchem"}, "error": None}

        def fake_verify_identifier(state, entity_id, field):
            ent = state.get_entity(entity_id)
            if ent is not None:
                ent.set_field_status(field, "verified", "lookup")
            return {"verified": True, "entity_id": entity_id, "field": field, "message": "ok"}

        monkeypatch.setattr(composites_mod, "lookup_compound", fake_lookup_compound)
        monkeypatch.setattr(composites_mod, "verify_identifier", fake_verify_identifier)

    def _by_type(self, engine: AgentEngine, type_name: str) -> list[Entity]:
        return [e for e in engine.state.list_entities() if e.type == type_name]

    def test_default_path_materializes_compounds_from_filenames(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        seen_prompts = self._fake_extract_plan_chat_model(monkeypatch)
        self._stub_compound_lookup(monkeypatch)

        engine = _engine(self._state_with_compound_files(tmp_path))

        # BEFORE: the default path has produced zero compounds.
        assert self._by_type(engine, "MolecularEntity") == []

        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        # The real leaf actually saw the compound-bearing filenames.
        joined = "\n".join(seen_prompts)
        assert "Silychristin+Verapamil" in joined
        assert "Diclofenac+BSP" in joined

        # AFTER: 0 → N MolecularEntities, one per filename-derived compound name.
        chems = self._by_type(engine, "MolecularEntity")
        assert {c.fields.get("name") for c in chems} == self._EXPECTED_NAMES
        assert result["compounds"] == len(self._EXPECTED_NAMES)

    def test_names_only_identifiers_come_from_the_lookup(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """D5: only the NAME reaches ``resolve_compound``; each MolecularEntity's
        CAS/CID is the LOOKED-UP value (never invented from the filename)."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._fake_extract_plan_chat_model(monkeypatch)
        self._stub_compound_lookup(monkeypatch)

        engine = _engine(self._state_with_compound_files(tmp_path))
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)

        chems = {c.fields.get("name"): c for c in self._by_type(engine, "MolecularEntity")}
        assert set(chems) == self._EXPECTED_NAMES, "all filename compounds minted"
        # The CAS on each entity is the name's LOOKED-UP value, not a fabrication.
        for name, expected in self._LOOKUP_BY_NAME.items():
            assert chems[name].fields.get("cas") == expected["cas"]


class TestPublicationFromPDF:
    """Issue #245 — when a plan publication's "title" is actually a PDF FILENAME,
    `_materialize_plan` must recover the real DOI/title from the PDF *text* and
    resolve with that, never query Crossref with the bare filename.

    The OATP1C1 symptom: ``extract_plan`` returned
    ``{'title': 'Wagenaars_etal_2025_OATP1C1.pdf'}`` and Crossref answered "no
    confident DOI match" — a title search on a filename can never match. These
    tests stub the PDF reader and the Crossref/DOI lookups so they are fully
    offline. Contract:

    * a filename-title backed by a PDF whose text contains a DOI → resolved by
      that DOI (``draft_publication_with_authors``), never by the filename;
    * a PDF text with a real title but no DOI → resolved by the extracted title;
    * neither DOI nor title recoverable → NO Crossref call with the filename;
      the publication is skipped/deferred gracefully.
    """

    _PDF_NAME = "Wagenaars_etal_2025_OATP1C1.pdf"

    def _enable_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")

    def _state_with_pdf(self, tmp_path: Path) -> CrateState:
        """A titled state whose scanned-file inventory includes a PDF under an
        approved root (the PDF need only EXIST on disk; its text is stubbed)."""
        pdf_path = tmp_path / self._PDF_NAME
        pdf_path.write_bytes(b"%PDF-1.4 stub")  # real bytes irrelevant — reader stubbed
        state = CrateState()
        state.metadata.title = "OATP1C1 transporter study"
        state.metadata.description = "An in vitro OATP1C1 uptake assay."
        state.approved_scan_roots.add(str(tmp_path.resolve()))
        state.scanned_files = [
            FileClassification(
                path=str(pdf_path),
                filename=self._PDF_NAME,
                size=pdf_path.stat().st_size,
                mime_type="application/pdf",
                first_rows=None,
            )
        ]
        return state

    def _stub_extract_plan_filename(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Make the leaf return the PDF FILENAME as the publication title (#245)."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        def fake_extract_plan(context, *, overrides=None, usage_sink=None):
            return {"publications": [{"title": self._PDF_NAME}]}

        monkeypatch.setattr(pipeline_mod, "extract_plan", fake_extract_plan)

    def _by_type(self, engine: AgentEngine, type_name: str) -> list[Entity]:
        return [e for e in engine.state.list_entities() if e.type == type_name]

    def test_pdf_doi_in_text_resolves_by_doi_not_filename(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """A PDF whose text carries a DOI → resolved by that DOI, not the filename."""
        import builder.agents.pipeline.pipeline as pipeline_mod
        import builder.tools.composites as composites_mod
        import builder.tools.scanner as scanner_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan_filename(monkeypatch)

        # The PDF text the reader returns, carrying a real DOI on the first page.
        def fake_extract_pdf_text(path):
            return (
                "[Page 1]\n"
                "[Text] Hepatic OATP1C1 mediates thyroid hormone uptake\n"
                "[Text] https://doi.org/10.1016/j.example.2025.01.002\n"
                "[Text] Fabian Wagenaars, et al. 2025\n"
            )

        monkeypatch.setattr(scanner_mod, "extract_pdf_text", fake_extract_pdf_text)

        # resolve_publication (title path) must NOT be called with the filename.
        def boom_search(query):  # pragma: no cover - must not run
            raise AssertionError(f"Crossref title search must not be called with {query!r}")

        monkeypatch.setattr(composites_mod, "search_works_by_title", boom_search)

        # The DOI path delegates to draft_publication_with_authors → lookup_doi.
        seen_dois: list[str] = []

        def fake_lookup_doi(doi):
            seen_dois.append(doi)
            return {
                "found": True,
                "data": {
                    "identifier": "10.1016/j.example.2025.01.002",
                    "name": "Hepatic OATP1C1 mediates thyroid hormone uptake",
                    "author": [{"givenName": "Fabian", "familyName": "Wagenaars"}],
                },
                "error": None,
            }

        monkeypatch.setattr(composites_mod, "lookup_doi", fake_lookup_doi)

        engine = _engine(self._state_with_pdf(tmp_path))
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        # The DOI extracted from the PDF text drove the resolution.
        assert seen_dois == ["10.1016/j.example.2025.01.002"]
        pubs = self._by_type(engine, "Publication")
        assert len(pubs) == 1
        assert pubs[0].fields.get("identifier") == "10.1016/j.example.2025.01.002"
        assert result["publications"] >= 1
        # The bare filename never lingers as a deferred title.
        assert self._PDF_NAME not in result["publications_deferred"]

    def test_pdf_title_no_doi_resolves_by_extracted_title(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """A PDF with a real title but no DOI → resolved by the extracted title."""
        import builder.agents.pipeline.pipeline as pipeline_mod
        import builder.tools.composites as composites_mod
        import builder.tools.scanner as scanner_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan_filename(monkeypatch)

        real_title = "Hepatic OATP1C1 mediates thyroid hormone uptake"

        def fake_extract_pdf_text(path):
            # First non-trivial line is the article title; no DOI anywhere.
            return f"[Page 1]\n[Text] {real_title}\n[Text] Fabian Wagenaars, et al. 2025\n"

        monkeypatch.setattr(scanner_mod, "extract_pdf_text", fake_extract_pdf_text)

        # The title path must be called with the REAL extracted title, not the file.
        seen_titles: list[str] = []

        def fake_search(query):
            seen_titles.append(query)
            return [{"doi": "10.1016/j.example.2025.01.002", "title": real_title, "score": 99.0}]

        def fake_lookup_doi(doi):
            return {
                "found": True,
                "data": {
                    "identifier": "10.1016/j.example.2025.01.002",
                    "name": real_title,
                    "author": [{"givenName": "Fabian", "familyName": "Wagenaars"}],
                },
                "error": None,
            }

        monkeypatch.setattr(composites_mod, "search_works_by_title", fake_search)
        monkeypatch.setattr(composites_mod, "lookup_doi", fake_lookup_doi)

        engine = _engine(self._state_with_pdf(tmp_path))
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        # Crossref was queried with the EXTRACTED TITLE, never the filename.
        assert seen_titles == [real_title]
        assert self._PDF_NAME not in seen_titles
        pubs = self._by_type(engine, "Publication")
        assert len(pubs) == 1
        assert result["publications"] >= 1
        assert self._PDF_NAME not in result["publications_deferred"]

    def test_pdf_no_doi_or_title_does_not_query_with_filename(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """No recoverable DOI/title → no Crossref-by-filename; graceful skip/defer."""
        import builder.agents.pipeline.pipeline as pipeline_mod
        import builder.tools.composites as composites_mod
        import builder.tools.scanner as scanner_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan_filename(monkeypatch)

        # The PDF text is unreadable/empty — nothing to recover.
        monkeypatch.setattr(scanner_mod, "extract_pdf_text", lambda path: None)

        def boom_search(query):  # pragma: no cover - must not run
            raise AssertionError(f"Crossref title search must not be called with {query!r}")

        def boom_lookup_doi(doi):  # pragma: no cover - must not run
            raise AssertionError("DOI lookup must not run without a recovered DOI")

        monkeypatch.setattr(composites_mod, "search_works_by_title", boom_search)
        monkeypatch.setattr(composites_mod, "lookup_doi", boom_lookup_doi)

        engine = _engine(self._state_with_pdf(tmp_path))
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        # Nothing minted; the filename did NOT reach Crossref. Skipped gracefully:
        # it is not counted as a resolved publication.
        assert self._by_type(engine, "Publication") == []
        assert result["publications"] == 0
        # The bare filename is never surfaced as a "deferred title" to retry with —
        # a filename can never become a confident title match.
        assert self._PDF_NAME not in result["publications_deferred"]


class TestTokenAccounting:
    """Issue #221 — the deterministic spine's leaf LLM calls must record their
    token usage to ``profile.ndjson`` in the SAME shape the ReAct model node
    uses, so ``eval/runner.py`` mines real per-case tokens for ``--arch pipeline``
    (it previously recorded 0 because the leaves' usage was discarded).

    Offline: the leaf is stubbed and reports a known usage payload via the
    ``usage_sink`` the spine passes it; the engine writes ``node_end``/
    ``node="model"`` events that :func:`eval.metrics.mine_profile_metrics` sums.
    """

    def _seeded_state(self) -> CrateState:
        import uuid

        state = CrateState()
        # A unique session id so each test writes its OWN profile.ndjson; the
        # _engine() default is second-precision, which collides under -n 2 and
        # would mix one test's model events into another's profile (read below).
        state.session_id = f"tok-{uuid.uuid4().hex[:12]}"
        state.metadata.title = "TPO inhibition dose-response screen"
        state.metadata.description = "A cell-based in vitro TPO inhibition assay."
        state.add_entity(_entity("inv1", "Investigation", name="Inv"))
        state.add_entity(_entity("st1", "Study", name="St", investigation_id="inv1"))
        state.add_entity(_entity("as1", "Assay", name="As", study_id="st1"))
        # Two bare entities (missing a description) for the leaf to enrich, so the
        # spine makes >1 leaf call and we assert usage is ACCUMULATED across them.
        state.add_entity(_entity("chem1", "MolecularEntity", name="Methimazole"))
        state.add_entity(_entity("per1", "Person", name="Jane Doe"))
        return state

    def _mine(self, engine: AgentEngine):
        import shutil
        from pathlib import Path

        from builder.tools.dashboard import read_profile
        from builder.tools.profiler import SESSION_DIR
        from eval.metrics import mine_profile_metrics

        engine.close_profiler()  # flush + close before reading
        session_dir = Path(SESSION_DIR) / engine.state.session_id
        try:
            return mine_profile_metrics(read_profile(session_dir / "profile.ndjson"))
        finally:
            # Don't litter sessions/ with this test's unique-id profile dir.
            shutil.rmtree(session_dir, ignore_errors=True)

    def test_draft_entities_records_accumulated_tokens(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")

        def fake_leaf(entity_type, context, *, overrides=None, usage_sink=None):
            # Each leaf call reports a known usage payload through the sink.
            if usage_sink is not None:
                usage_sink(100, 20, "gpt-4o-mini")
            return {"description": f"drafted {entity_type}"}

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", fake_leaf)

        engine = _engine(self._seeded_state())
        totals = {"input_tokens": 0, "output_tokens": 0}
        sink = pipeline_mod._make_usage_logger(engine, totals)
        result = pipeline_mod._draft_entities(engine, sink)

        # The spine enriched every entity missing a descriptive field; at least
        # the two bare domain entities, so >1 leaf call is made (we assert usage
        # is ACCUMULATED, not just recorded once).
        n_calls = len(result["drafted"])
        assert n_calls >= 2

        # The running accumulator summed every leaf call (n_calls × 100/20).
        assert totals["input_tokens"] == 100 * n_calls
        assert totals["output_tokens"] == 20 * n_calls

        # …and those landed in profile.ndjson as node_end/model events, so the
        # eval runner mines them identically to the ReAct arm.
        pm = self._mine(engine)
        assert pm.input_tokens == 100 * n_calls
        assert pm.output_tokens == 20 * n_calls
        assert pm.total_tokens == 120 * n_calls

    def test_no_provider_records_clean_zero(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        # No provider → the leaf is never called and no model events are written.
        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: None)

        def boom(*args, **kwargs):  # pragma: no cover - must not run
            raise AssertionError("leaf must not run without a provider")

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", boom)

        engine = _engine(self._seeded_state())
        pipeline_mod._draft_entities(engine)

        pm = self._mine(engine)
        assert pm.input_tokens == 0
        assert pm.output_tokens == 0
        assert pm.total_tokens == 0

    def test_run_pipeline_returns_usage_dict(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """``run_pipeline``'s result additively surfaces the accumulated usage."""
        import builder.agents.pipeline.pipeline as pipeline_mod
        from builder.agents.pipeline.pipeline import run_pipeline

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")
        # Make the plan stage a no-op (empty plan) so only the drafter leaf runs.
        monkeypatch.setattr(pipeline_mod, "extract_plan", lambda *a, **k: {})

        def fake_leaf(entity_type, context, *, overrides=None, usage_sink=None):
            if usage_sink is not None:
                usage_sink(50, 10, "gpt-4o-mini")
            return {"description": f"drafted {entity_type}"}

        monkeypatch.setattr(pipeline_mod, "draft_entity_fields", fake_leaf)

        import shutil
        from pathlib import Path

        from builder.tools.profiler import SESSION_DIR

        engine = _engine(self._seeded_state())
        try:
            result = run_pipeline(engine)
        finally:
            engine.close_profiler()
            shutil.rmtree(Path(SESSION_DIR) / engine.state.session_id, ignore_errors=True)

        assert "usage" in result
        usage = result["usage"]
        # At least the two seeded bare entities were drafted (50/10 each).
        assert usage["input_tokens"] >= 100
        assert usage["output_tokens"] >= 20
        assert usage["total_tokens"] == usage["input_tokens"] + usage["output_tokens"]


class TestDeterminism:
    def test_identical_graph_hash_across_runs(self) -> None:
        """Same input ⇒ identical built @graph hash — the headline win to assert."""
        from builder.agents.pipeline.pipeline import run_pipeline
        from eval.metrics import crate_graph_hash

        e1 = _engine()
        run_pipeline(e1)
        h1 = crate_graph_hash(e1.state)

        e2 = _engine()
        run_pipeline(e2)
        h2 = crate_graph_hash(e2.state)

        assert h1 == h2

    def test_determinism_holds_with_seeded_entities(self) -> None:
        """Determinism holds even when state carries non-backbone entities."""
        from builder.agents.pipeline.pipeline import run_pipeline
        from eval.metrics import crate_graph_hash

        def seeded() -> CrateState:
            state = CrateState()
            state.metadata.title = "Deterministic crate"
            state.add_entity(_entity("chem1", "MolecularEntity", name="Triiodothyronine"))
            return state

        e1 = _engine(seeded())
        run_pipeline(e1)
        e2 = _engine(seeded())
        run_pipeline(e2)
        assert crate_graph_hash(e1.state) == crate_graph_hash(e2.state)


class TestGatherContext:
    """`_gather_context` now reads non-tabular rich file BODIES (Issue #231).

    Before #231 the spine's single bounded extraction leaf only ever saw
    filenames plus tiny tabular ``first_rows`` previews, so ``.json`` / ``.docx``
    / ``.pdf`` rich files contributed nothing and ``extract_plan`` returned an
    empty plan (the backbone then fell back to literal default names). These
    tests pin the fix: bodies of readable non-tabular files appear in the
    gathered context, bounded by ``_MAX_CONTEXT_CHARS`` (per-file + total),
    confined to ``approved_scan_roots``, and never raising out of the spine. The
    two strict no-op gates (empty context ⇒ ``""``) are preserved.
    """

    def _write_json(self, root: Path, marker: str) -> FileClassification:
        path = root / "S-VHPS26.json"
        path.write_text(
            f'{{"studyTitle": "{marker}", "organism": "Rattus norvegicus"}}',
            encoding="utf-8",
        )
        return FileClassification(
            path=str(path),
            filename=path.name,
            size=path.stat().st_size,
            mime_type="application/json",
            first_rows=None,
        )

    def _write_docx(self, root: Path, marker: str) -> FileClassification:
        from docx import Document  # type: ignore[import-untyped]

        path = root / "SOP.docx"
        doc = Document()
        doc.add_paragraph(marker)
        doc.add_paragraph("Standard operating procedure for the assay.")
        doc.save(str(path))
        return FileClassification(
            path=str(path),
            filename=path.name,
            size=path.stat().st_size,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            first_rows=None,
        )

    def test_reads_json_and_docx_bodies_into_context(self, tmp_path: Path) -> None:
        """Body substrings from non-tabular files under an approved root appear."""
        from builder.agents.pipeline.pipeline import _gather_context

        json_fc = self._write_json(tmp_path, "JSONBODYMARKER")
        docx_fc = self._write_docx(tmp_path, "DOCXBODYMARKER")

        state = CrateState()
        state.approved_scan_roots.add(str(tmp_path.resolve()))
        state.scanned_files = [json_fc, docx_fc]
        engine = _engine(state)

        context = _gather_context(engine)

        # The bodies — not just the filenames — reached the context.
        assert "JSONBODYMARKER" in context
        assert "DOCXBODYMARKER" in context
        # Filenames are still listed.
        assert "S-VHPS26.json" in context
        assert "SOP.docx" in context

    def test_prefers_cheap_first_rows_when_present(self, tmp_path: Path) -> None:
        """A file carrying a tabular preview uses it; disk is not re-read for it."""
        from builder.agents.pipeline.pipeline import _gather_context
        from builder.state import FileClassification

        # A file that DOES carry first_rows — the cheap preview must be used and the
        # body reader must not be invoked for it (the path need not even exist).
        tabular = FileClassification(
            path=str(tmp_path / "missing.csv"),
            filename="data.csv",
            size=10,
            mime_type="text/csv",
            first_rows=["col_a,col_b", "1,2"],
        )
        state = CrateState()
        state.approved_scan_roots.add(str(tmp_path.resolve()))
        state.scanned_files = [tabular]
        engine = _engine(state)

        context = _gather_context(engine)
        assert "col_a,col_b" in context

    def test_reads_are_confined_to_approved_scan_roots(self, tmp_path: Path) -> None:
        """A body OUTSIDE every approved root is never read into the context."""
        from builder.agents.pipeline.pipeline import _gather_context

        approved = tmp_path / "approved"
        approved.mkdir()
        outside = tmp_path / "outside"
        outside.mkdir()

        inside_fc = self._write_json(approved, "INSIDEMARKER")
        outside_fc = self._write_json(outside, "OUTSIDEMARKER")
        # Re-point filename so the two are distinguishable in the listing.
        outside_fc.filename = "outside.json"

        state = CrateState()
        state.approved_scan_roots.add(str(approved.resolve()))
        state.scanned_files = [inside_fc, outside_fc]
        engine = _engine(state)

        context = _gather_context(engine)
        assert "INSIDEMARKER" in context
        # The out-of-root body is fail-closed: its content never appears.
        assert "OUTSIDEMARKER" not in context

    def test_per_file_and_total_budget_are_bounded(self, tmp_path: Path) -> None:
        """Per-file and total context are capped by `_MAX_CONTEXT_CHARS`."""
        import builder.agents.pipeline.pipeline as pipeline_mod
        from builder.agents.pipeline.pipeline import _gather_context
        from builder.state import FileClassification

        cap = pipeline_mod._MAX_CONTEXT_CHARS
        assert isinstance(cap, int) and cap > 0

        state = CrateState()
        state.approved_scan_roots.add(str(tmp_path.resolve()))
        files: list[FileClassification] = []
        # Several large files, each far bigger than the cap, so both the per-file
        # and the total budget must clamp the result.
        for i in range(6):
            path = tmp_path / f"big{i}.txt"
            path.write_text("X" * (cap * 4), encoding="utf-8")
            files.append(
                FileClassification(
                    path=str(path),
                    filename=path.name,
                    size=path.stat().st_size,
                    mime_type="text/plain",
                    first_rows=None,
                )
            )
        state.scanned_files = files
        engine = _engine(state)

        context = _gather_context(engine)
        # Total context stays within a small multiple of the documented cap (the
        # total budget is the binding ceiling — not 6x the per-file body).
        assert len(context) <= cap * 3

    def test_no_readable_files_is_strict_noop(self, tmp_path: Path) -> None:
        """Nothing readable ⇒ ``""`` so the no-provider determinism gate holds."""
        from builder.agents.pipeline.pipeline import _gather_context
        from builder.state import FileClassification

        # A binary file with no first_rows whose body reader returns None, under an
        # approved root, on an untitled/undescribed crate ⇒ no usable context.
        path = tmp_path / "blob.pzfx"
        path.write_bytes(b"\x00\x01\x02\x03binary-not-text\x00")
        state = CrateState()
        state.approved_scan_roots.add(str(tmp_path.resolve()))
        state.scanned_files = [
            FileClassification(
                path=str(path),
                filename=path.name,
                size=path.stat().st_size,
                mime_type="application/octet-stream",
                first_rows=None,
            )
        ]
        engine = _engine(state)

        # Filenames alone still list, but no BODY content — and an untitled crate
        # with only a filename listing is still usable context for the listing
        # path. The strict no-op we must preserve is the *fully empty* one:
        state_empty = CrateState()
        engine_empty = _engine(state_empty)
        assert _gather_context(engine_empty) == ""

        # The binary body itself contributed nothing readable.
        context = _gather_context(engine)
        assert "binary-not-text" not in context


# ---------------------------------------------------------------------------
# Issue #232 (b) — deterministic person-name splitting must not mis-place a
# surname into givenName, and EVERY draft_person path must be ISA-shaped.
# Appended at the END to avoid colliding with the sibling (#231) edits.
# ---------------------------------------------------------------------------


class TestSplitPersonName:
    """`_split_person_name` contract.

    - "Last, First" / comma form is inverted to (given, family).
    - Trailing punctuation around the comma is stripped.
    - A lone bare surname is treated as a family-name candidate, NOT silently
      mis-placed into givenName (the "Wagenaars" bug).
    - A plain "First Last" still splits on the last token.
    """

    def test_comma_form_is_inverted(self) -> None:
        from builder.agents.pipeline.pipeline import _split_person_name

        assert _split_person_name("Wagenaars, J.") == ("J.", "Wagenaars")

    def test_comma_form_full_given(self) -> None:
        from builder.agents.pipeline.pipeline import _split_person_name

        assert _split_person_name("Lovelace, Ada") == ("Ada", "Lovelace")

    def test_comma_form_multi_token_given(self) -> None:
        from builder.agents.pipeline.pipeline import _split_person_name

        assert _split_person_name("van Helsing, Abraham A.") == (
            "Abraham A.",
            "van Helsing",
        )

    def test_lone_surname_is_family_candidate_not_given(self) -> None:
        from builder.agents.pipeline.pipeline import _split_person_name

        given, family = _split_person_name("Wagenaars")
        # The bug: the surname used to land in givenName. It must not.
        assert given == ""
        assert family == "Wagenaars"

    def test_plain_first_last_unchanged(self) -> None:
        from builder.agents.pipeline.pipeline import _split_person_name

        assert _split_person_name("Ada Lovelace") == ("Ada", "Lovelace")

    def test_three_token_name(self) -> None:
        from builder.agents.pipeline.pipeline import _split_person_name

        assert _split_person_name("Ada King Lovelace") == ("Ada King", "Lovelace")

    def test_empty_name(self) -> None:
        from builder.agents.pipeline.pipeline import _split_person_name

        assert _split_person_name("   ") == ("", "")


class TestDraftPersonSplit:
    """Every `draft_person` path must be able to produce a split, ISA-shaped
    Person — not only the materialize path.

    `draft_person` falls back to `_split_person_name` when the caller supplies
    neither `givenName` nor `familyName`, so a bare `draft_person(name=...)` call
    (ReAct / guidance / direct) is ISA-conformant. An explicit split the caller
    passes wins. ORCID stays empty (D5)."""

    def test_falls_back_to_split_when_no_hint(self) -> None:
        from builder.state import CrateState
        from builder.tools.drafters import draft_person

        state = CrateState()
        person = draft_person(state, "Ada Lovelace", {})
        assert person.fields.get("givenName") == "Ada"
        assert person.fields.get("familyName") == "Lovelace"
        assert person.fields.get("name") == "Ada Lovelace"
        assert not person.fields.get("orcid")

    def test_comma_form_via_draft_person(self) -> None:
        from builder.state import CrateState
        from builder.tools.drafters import draft_person

        state = CrateState()
        person = draft_person(state, "Wagenaars, J.", {})
        assert person.fields.get("givenName") == "J."
        assert person.fields.get("familyName") == "Wagenaars"

    def test_explicit_split_hint_wins(self) -> None:
        from builder.state import CrateState
        from builder.tools.drafters import draft_person

        state = CrateState()
        person = draft_person(state, "Ada Lovelace", {"givenName": "A.", "familyName": "Lovelace"})
        # The caller's explicit split is preserved, not recomputed from name.
        assert person.fields.get("givenName") == "A."
        assert person.fields.get("familyName") == "Lovelace"

    def test_partial_hint_does_not_trigger_split(self) -> None:
        """If the caller supplies ONLY givenName, the split fallback does not run
        (we don't second-guess a partial explicit hint)."""
        from builder.state import CrateState
        from builder.tools.drafters import draft_person

        state = CrateState()
        person = draft_person(state, "Ada Lovelace", {"givenName": "Ada"})
        assert person.fields.get("givenName") == "Ada"
        # familyName is left to a later step; the fallback did not overwrite given.
        assert not person.fields.get("familyName")


# ---------------------------------------------------------------------------
# Issue #232 (a) — the plan's backbone name must land on the already-scaffolded
# backbone (the scaffold runs BEFORE the plan is materialized).
# ---------------------------------------------------------------------------


class TestMaterializeBackboneNaming:
    """`_materialize_plan` must merge the plan's Study name/description onto the
    backbone that `_scaffold_backbone` already created.

    `scaffold_isa_backbone` *reuses* an existing Study (its hints reach the
    drafter only on creation). So on an UNTITLED crate the Study keeps the generic
    "Study" placeholder even when the plan supplies a real name — this asserts the
    merge now lands (fill-don't-clobber).
    """

    def _enable_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")

    def _stub_extract_plan(self, monkeypatch: pytest.MonkeyPatch, plan: dict) -> None:
        import builder.agents.pipeline.pipeline as pipeline_mod

        def fake_extract_plan(context, *, overrides=None, usage_sink=None):
            return dict(plan)

        monkeypatch.setattr(pipeline_mod, "extract_plan", fake_extract_plan)

    def _study(self, engine: AgentEngine) -> Entity:
        return next(e for e in engine.state.list_entities() if e.type == "Study")

    def _untitled_with_context(self) -> CrateState:
        """An UNTITLED crate that still carries usable context (a scanned file), so
        `_materialize_plan` does not no-op on the context gate but the backbone is
        named with the generic default rather than a title."""
        from builder.state import FileClassification

        state = CrateState()
        state.scanned_files = [
            FileClassification(
                path="data/notes.txt",
                filename="notes.txt",
                size=10,
                mime_type="text/plain",
                first_rows=None,
            )
        ]
        return state

    def test_plan_study_name_lands_on_scaffolded_backbone(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """On an untitled crate, the plan's study.name must overwrite the generic
        "Study" placeholder the scaffold left behind."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(
            monkeypatch,
            {"study": {"name": "TPO inhibition study", "description": "A TPO assay."}},
        )

        engine = _engine(self._untitled_with_context())
        pipeline_mod._scaffold_backbone(engine)
        # Sanity: the scaffold left the generic default name (the pre-fix state).
        assert self._study(engine).fields.get("name") == "Study"

        result = pipeline_mod._materialize_plan(engine)

        study = self._study(engine)
        assert study.fields.get("name") == "TPO inhibition study"
        assert study.fields.get("description") == "A TPO assay."
        assert result["study"] == 1

    def test_no_plan_name_keeps_default_nonempty(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Regression: with no plan study (and no title) the Study keeps a
        non-empty default name — ISA REQUIRES a non-empty Study name."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch, {"compounds": []})

        engine = _engine(self._untitled_with_context())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)

        name = self._study(engine).fields.get("name")
        assert isinstance(name, str) and name.strip(), "Study name must stay non-empty"

    def test_existing_specific_name_is_not_clobbered(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Fill-don't-clobber: a Study that already carries a real (non-default)
        name must NOT be overwritten by the plan."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._enable_provider(monkeypatch)
        self._stub_extract_plan(monkeypatch, {"study": {"name": "Plan-supplied name"}})

        state = self._untitled_with_context()
        state.metadata.title = "A real, specific study title"
        engine = _engine(state)
        pipeline_mod._scaffold_backbone(engine)
        # The titled scaffold named the Study from the title (a real name).
        assert self._study(engine).fields.get("name") == "A real, specific study title"

        pipeline_mod._materialize_plan(engine)

        # The specific name wins over the plan (fill-don't-clobber).
        assert self._study(engine).fields.get("name") == "A real, specific study title"


# ---------------------------------------------------------------------------
# Issues #241 / #242 — pipeline progress + state persistence.
#
# #242: run_pipeline NEVER persisted CrateState, so a concurrent --dashboard read
# "No CrateState data available" and never live-updated. The spine must
# save_session at each phase boundary so the watched crate_state.json appears
# and changes as the build progresses.
#
# #241: the pipeline emitted NO progress, so the ~tens-of-seconds deterministic
# spine looked frozen. The spine must emit one concise line per phase through an
# injected progress callback, defaulting to a strict no-op so eval + determinism
# stay clean.
# ---------------------------------------------------------------------------


class TestPipelineStatePersistence:
    """#242 — run_pipeline persists CrateState at phase boundaries."""

    def _isolate_sessions(self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> Path:
        """Point save_session at a tmp sessions dir and reset its dedup cache."""
        import builder.tools.session as sess_mod

        sessions = tmp_path / "sessions"
        monkeypatch.setattr(sess_mod, "SESSION_DIR", sessions)
        monkeypatch.setattr(sess_mod, "_last_saved_state_hash", None)
        return sessions

    def test_pipeline_writes_crate_state_json(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """After run_pipeline, sessions/<id>/crate_state.json exists with entities."""
        from builder.agents.pipeline.pipeline import run_pipeline
        from builder.tools.session import load_session

        sessions = self._isolate_sessions(monkeypatch, tmp_path)
        engine = _engine()
        run_pipeline(engine)

        state_path = sessions / engine.state.session_id / "crate_state.json"
        assert state_path.is_file(), "run_pipeline must persist crate_state.json (#242)"

        loaded = load_session(engine.state.session_id)
        assert loaded is not None
        types = {e.type for e in loaded.list_entities()}
        # The scaffolded backbone is persisted, so a dashboard sees real entities.
        assert {"Investigation", "Study", "Assay"} <= types

    def test_pipeline_saves_at_least_once(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """A save callback is invoked during the spine so a concurrent dashboard
        sees progress, not just the final state."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._isolate_sessions(monkeypatch, tmp_path)
        saves: list[str] = []

        def fake_save(state, *, always_write: bool = False, **kw):
            saves.append(state.session_id)
            return {"success": True, "session_id": state.session_id, "skipped": False}

        engine = _engine()
        pipeline_mod.run_pipeline(engine, save=fake_save)

        # At least one save during the spine (incremental dashboard updates).
        assert len(saves) >= 1

    def test_pipeline_saves_at_multiple_phase_boundaries(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        """Multiple saves at phase boundaries (scaffold + each validate) so the
        watched file changes incrementally during the run."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._isolate_sessions(monkeypatch, tmp_path)
        saves: list[int] = []

        def fake_save(state, *, always_write: bool = False, **kw):
            saves.append(len(state.list_entities()))
            return {"success": True, "session_id": state.session_id, "skipped": False}

        engine = _engine()
        pipeline_mod.run_pipeline(engine, save=fake_save)

        # Scaffold boundary + fix-loop validate boundary => more than one save.
        assert len(saves) >= 2


class TestPipelineProgress:
    """#241 — run_pipeline emits one concise progress line per phase."""

    def test_progress_callback_receives_phase_lines(self) -> None:
        from builder.agents.pipeline.pipeline import run_pipeline

        lines: list[str] = []
        engine = _engine()
        run_pipeline(engine, progress=lines.append)

        joined = "\n".join(lines).lower()
        # Each major phase surfaces a concise line.
        assert "scaffold" in joined
        assert "validat" in joined
        # The lines are short, human-readable phase markers (not raw dicts).
        assert lines
        assert all(len(line) < 200 for line in lines)

    def test_progress_defaults_to_noop(self) -> None:
        """With no progress callback the spine emits nothing (clean eval/tests)."""
        from builder.agents.pipeline.pipeline import run_pipeline

        engine = _engine()
        # Must not raise and must not print — a missing callback is a strict no-op.
        # This is about the CALLBACK, not conformance: an empty state has no
        # process parameters to state, so `ok` is legitimately False here (see
        # test_scaffold_only_reaches_base_and_isa).
        result = run_pipeline(engine)
        assert set(result["conformance"]) == {"base", "isa", "tox"}


# ---------------------------------------------------------------------------
# Issue #262 — the pipeline must DETERMINISTICALLY materialize the standard
# in-vitro process chain AND attach the scanned data files. Both run with NO
# provider configured (they are pure, code-driven steps), so the crate is never
# structurally hollow (0 lab_processes / 0 files) even on the offline spine.
# ---------------------------------------------------------------------------


def _scanned(
    path: str, filename: str, mime: str = "application/octet-stream"
) -> FileClassification:
    return FileClassification(
        path=path, filename=filename, size=64, mime_type=mime, first_rows=None
    )


class TestMaterializeStandardProcessChain:
    """`_materialize_plan` deterministically drafts the standard 4-step in-vitro
    chain (CellCulture → Exposure → EndpointReadout → DataAnalysis) under the
    scaffolded Assay — regardless of whether an LLM provider is configured.

    Before #262 the pipeline only drafted a chain from a provider-extracted plan,
    so the no-provider crate was structurally hollow (``lab_processes: []``).
    """

    def _by_type(self, engine: AgentEngine, type_name: str) -> list[Entity]:
        return [e for e in engine.state.list_entities() if e.type == type_name]

    def _titled_state(self) -> CrateState:
        state = CrateState()
        state.metadata.title = "TPO inhibition dose-response screen"
        return state

    def _no_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Force the no-provider path so the plan-driven section is a strict no-op."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: None)

        def boom(*args, **kwargs):  # pragma: no cover - must not run without a provider
            raise AssertionError("extract_plan must not run without a provider")

        monkeypatch.setattr(pipeline_mod, "extract_plan", boom)

    def test_no_provider_drafts_full_chain_under_assay(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """With NO provider, the standard 4-step chain is wired under the Assay."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._no_provider(monkeypatch)

        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        procs = self._by_type(engine, "LabProcess")
        assert {p.fields.get("process_type") for p in procs} == {
            "CellCulture",
            "Exposure",
            "EndpointReadout",
            "DataAnalysis",
        }
        # Every process belongs to the scaffolded Assay.
        assay_id = next(e.entity_id for e in engine.state.list_entities() if e.type == "Assay")
        for proc in procs:
            assay_ref = proc.fields.get("assay") or proc.fields.get("partOf")
            ref = assay_ref.get("@id") if isinstance(assay_ref, dict) else assay_ref
            if ref is not None:
                assert str(ref).lstrip("#") == assay_id
        # The result trace reflects the materialized chain.
        assert result["processes"] >= 4

    def test_endpoint_and_analysis_carry_outputs(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """§14.3: EndpointReadout/DataAnalysis must carry a result (no Violation trap).

        ``draft_process_chain`` synthesizes placeholder File outputs for the two
        data-producing steps that have no build-time fallback, so neither dangles.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._no_provider(monkeypatch)
        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)

        procs = self._by_type(engine, "LabProcess")
        for ptype in ("EndpointReadout", "DataAnalysis"):
            proc = next(p for p in procs if p.fields.get("process_type") == ptype)
            assert proc.fields.get("result"), (
                f"{ptype} must carry a result output (the §14.3 no-output trap)"
            )

    def test_chain_is_idempotent_across_repeats(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Re-running the spine mints no duplicate processes (deterministic ids)."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._no_provider(monkeypatch)
        engine = _engine(self._titled_state())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)
        n1 = len(self._by_type(engine, "LabProcess"))
        pipeline_mod._materialize_plan(engine)
        n2 = len(self._by_type(engine, "LabProcess"))
        assert n1 == n2 == 4

    def test_chain_present_after_full_run_pipeline_no_provider(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """The full offline ``run_pipeline`` lands a non-empty process chain."""
        from builder.agents.pipeline.pipeline import run_pipeline

        self._no_provider(monkeypatch)
        engine = _engine(self._titled_state())
        run_pipeline(engine)
        procs = self._by_type(engine, "LabProcess")
        assert {p.fields.get("process_type") for p in procs} == {
            "CellCulture",
            "Exposure",
            "EndpointReadout",
            "DataAnalysis",
        }


class TestMaterializeAttachScannedFiles:
    """`_materialize_plan` deterministically adds every scanned data file as a File
    entity and links it to a process/assay — regardless of provider (#262).

    Before #262 the scanned files were never added to the crate (``files: []``);
    the pipeline now attaches them through the existing ``attach_files`` composite.
    """

    def _by_type(self, engine: AgentEngine, type_name: str) -> list[Entity]:
        return [e for e in engine.state.list_entities() if e.type == type_name]

    def _no_provider(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Force the no-provider path so file attachment is the only file source."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: None)

        def boom(*args, **kwargs):  # pragma: no cover - must not run without a provider
            raise AssertionError("extract_plan must not run without a provider")

        monkeypatch.setattr(pipeline_mod, "extract_plan", boom)

    def _state_with_files(self) -> CrateState:
        state = CrateState()
        state.metadata.title = "TPO inhibition dose-response screen"
        state.scanned_files = [
            _scanned(
                "data/dose_response_raw.xlsx",
                "dose_response_raw.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ),
            _scanned("analysis/ic50_results.prism", "ic50_results.prism"),
            _scanned("docs/SOP.pdf", "SOP.pdf", "application/pdf"),
        ]
        return state

    def test_no_provider_attaches_every_scanned_file(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """With NO provider, each scanned file becomes a File entity in the crate."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._no_provider(monkeypatch)
        engine = _engine(self._state_with_files())
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        files = self._by_type(engine, "File")
        attached_names = {f.fields.get("name") for f in files}
        assert {
            "dose_response_raw.xlsx",
            "ic50_results.prism",
            "SOP.pdf",
        } <= attached_names
        # The result trace records how many were attached (exactly the 3 scanned).
        assert result["files"] == 3

    def test_attached_files_are_linked_to_a_process_or_assay(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Every scanned File is referenced by a process (result/object) or under
        the Assay's hasPart — no silently orphaned data file."""
        import builder.agents.pipeline.pipeline as pipeline_mod
        from builder.tools.provenance import _ref_ids

        self._no_provider(monkeypatch)
        engine = _engine(self._state_with_files())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)

        # Collect every entity_id referenced by a process input/output or any
        # entity's hasPart (the two ways a File is wired into the structure).
        referenced: set[str] = set()
        for ent in engine.state.list_entities():
            for fld in (
                "object",
                "result",
                "input",
                "output",
                "samples",
                "hasPart",
                "has_part",
            ):
                referenced |= _ref_ids(ent.fields.get(fld))

        scanned_basenames = {
            "dose_response_raw.xlsx",
            "ic50_results.prism",
            "SOP.pdf",
        }
        scanned_files = [
            f for f in self._by_type(engine, "File") if f.fields.get("name") in scanned_basenames
        ]
        assert scanned_files, "the scanned data files must be added as File entities"
        for fe in scanned_files:
            assert fe.entity_id in referenced, (
                f"scanned file {fe.entity_id} is orphaned — not linked to any "
                f"process/assay (hasPart/result/object)"
            )

    def test_no_scanned_files_attaches_nothing(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """With no scanned files the attach step attaches nothing (no crash).

        The chain composite still synthesizes its §14.3 placeholder File outputs,
        so the crate is not File-free; what must be true is that the *attachment*
        step adds zero scanned files (``result["files"] == 0``).
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._no_provider(monkeypatch)
        state = CrateState()
        state.metadata.title = "A titled but file-less crate"
        engine = _engine(state)
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)

        assert result["files"] == 0

    def test_attachment_is_idempotent(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Re-running mints no duplicate File entities (deduped by on-disk source)."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        self._no_provider(monkeypatch)
        engine = _engine(self._state_with_files())
        pipeline_mod._scaffold_backbone(engine)
        pipeline_mod._materialize_plan(engine)
        n1 = len(self._by_type(engine, "File"))
        pipeline_mod._materialize_plan(engine)
        n2 = len(self._by_type(engine, "File"))
        # The 3 scanned files plus any synthesized chain placeholders, stable.
        assert n1 == n2


class TestGatherContextPreviewedFilesGetABudget:
    """A previewed file must get a real budget slice, not a hard 3 rows (#378).

    `TestGatherContextMetadataFirst` sets ``first_rows=None`` on every file with
    the comment "forces the body-read path", so the entire class is structurally
    blind to the preview branch. These tests drive that branch specifically.

    Every file here points at a path that does NOT exist, so `_read_body_excerpt`
    fails closed and the emitted slice comes from ``first_rows`` — the fallback
    the fix must keep, and must stop treating as a 3-row ceiling.
    """

    _SENTINEL = "ROW{:03d}-marker"

    def _state(self, tmp_path: Path, spec: list[tuple[str, int]], row_pad: int = 0) -> CrateState:
        """Build a state whose files carry *n* sentinel preview rows each.

        ``row_pad`` widens each row, so a test can breach the total budget with
        only the first three rows — the bytes today's code emits outside the
        arithmetic.
        """
        state = CrateState()
        state.approved_scan_roots = {str(tmp_path)}
        for filename, n_rows in spec:
            state.scanned_files.append(
                FileClassification(
                    path=str(tmp_path / filename),  # deliberately never created
                    filename=filename,
                    size=1,
                    mime_type="text/csv",
                    first_rows=[self._SENTINEL.format(i) + "x" * row_pad for i in range(n_rows)],
                )
            )
        return state

    @staticmethod
    def _slice_for(context: str, filename: str) -> str:
        """The CONTENT `_gather_context` emitted for one file.

        The ``- <filename>`` header is stripped deliberately: leaving it in makes
        the length comparison below pass on filename length alone, which is a
        tautology rather than a budget assertion.
        """
        for block in context.split("\n- "):
            block = block.lstrip("- ")
            if block.startswith(filename):
                return block[len(filename) :].lstrip(":").strip()
        return ""

    def test_previewed_metadata_file_is_not_capped_at_three_rows(self, tmp_path):
        """The scanner already paid to read the preview; emit more than 3 rows.

        Non-tautological: the sentinel comes from `FileClassification.first_rows`
        handed to the real `_gather_context`, not from a hand-built string.
        """
        engine = _engine(self._state(tmp_path, [("assay_metadata.csv", 40)]))

        import builder.agents.pipeline.pipeline as pipeline_mod

        context = pipeline_mod._gather_context(engine)

        assert self._SENTINEL.format(0) in context
        assert self._SENTINEL.format(30) in context, "preview still truncated to 3 rows"

    def test_previewed_content_is_charged_against_the_total_budget(self, tmp_path):
        """Preview bytes must count against the ceiling, not bypass it.

        Today only the body path decrements `body_budget`, so previews are spent
        outside the arithmetic and the documented total is silently exceeded.
        """
        # Only 8 rows, but each is 5,000 chars wide: today's `first_rows[:3]`
        # slice alone emits ~30,000 chars per file, none of it charged.
        engine = _engine(
            self._state(
                tmp_path,
                [("assay_metadata.csv", 8), ("bulk_data.csv", 8)],
                row_pad=5000,
            )
        )

        import builder.agents.pipeline.pipeline as pipeline_mod
        from builder.agents.pipeline.pipeline import _MAX_CONTEXT_CHARS

        context = pipeline_mod._gather_context(engine)

        assert len(context) <= _MAX_CONTEXT_CHARS + 2000

    def test_priority_zero_outranks_priority_three_in_chars_not_just_order(self, tmp_path):
        """Metadata-first must mean CHARS, not merely position.

        This is the assertion today's code inverts: the priority-0 workbook gets
        298 chars while a priority-3 bulk file gets 2,049.
        """
        engine = _engine(
            self._state(tmp_path, [("assay_metadata.csv", 4000), ("bulk_data.csv", 4000)])
        )

        import builder.agents.pipeline.pipeline as pipeline_mod

        context = pipeline_mod._gather_context(engine)

        meta = self._slice_for(context, "assay_metadata.csv")
        bulk = self._slice_for(context, "bulk_data.csv")
        assert len(meta) > len(bulk), f"metadata {len(meta)} chars, bulk {len(bulk)} chars"

    def test_absent_metadata_file_yields_no_metadata_content(self, tmp_path):
        """HONESTY CONTROL — the assertions above come from the file, not scaffolding."""
        engine = _engine(self._state(tmp_path, [("bulk_data.csv", 40)]))

        import builder.agents.pipeline.pipeline as pipeline_mod

        context = pipeline_mod._gather_context(engine)

        assert "assay_metadata.csv" not in context
        assert self._SENTINEL.format(0) in context  # the bulk file still speaks


class TestPlanChainParameterOverlay:
    """Plan-stated experimental parameters must reach the process hints (#379).

    `_merge_plan_chain_names` read only `step["name"]`, so a model that volunteered
    `{"duration": "30 minutes"}` had it silently discarded — and every default-arm
    crate shipped 11 ontology-typed PropertyValues asserting `"unknown"`, `"NA"`
    and `"Standard medium"` that nobody stated.
    """

    @staticmethod
    def _hints_for(chain: list[dict], ptype: str) -> dict:
        return next(s["hints"] for s in chain if s["process_type"] == ptype)

    def test_plan_parameters_are_overlaid_onto_the_step_hints(self):
        import builder.agents.pipeline.pipeline as pipeline_mod

        chain = pipeline_mod._merge_plan_chain_names(
            {
                "process_chain": [
                    {
                        "process_type": "Exposure",
                        "name": "Dose",
                        "parameters": {"duration": "30 minutes", "microplate": "96-well"},
                    }
                ]
            }
        )

        assert self._hints_for(chain, "Exposure") == {
            "name": "Dose",
            "duration": "30 minutes",
            "microplate": "96-well",
        }

    def test_non_whitelisted_step_keys_are_not_overlaid(self):
        """HONESTY CONTROL — a whitelist, not a splat.

        This is the test that fails if someone "fixes" the red above by merging
        the step dict wholesale. Plan items are `additionalProperties: True`, so
        `object_hint` would become a LabProcess state field and an `entity_id`
        would hijack the process `@id` via `drafters._make_entity_id`.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        chain = pipeline_mod._merge_plan_chain_names(
            {
                "process_chain": [
                    {
                        "process_type": "Exposure",
                        "name": "Dose",
                        "object_hint": "cells",
                        "entity_id": "proc_hijack",
                        "parameters": {
                            "duration": "30 minutes",
                            "entity_id": "proc_hijack",
                            "units": "min",
                            "cas": "51-48-9",
                            "object_hint": "cells",
                        },
                    }
                ]
            }
        )

        hints = self._hints_for(chain, "Exposure")
        assert hints == {"name": "Dose", "duration": "30 minutes"}
        for smuggled in ("entity_id", "units", "cas", "object_hint"):
            assert smuggled not in hints

    def test_empty_parameter_values_do_not_overlay(self):
        """A present-but-blank value must not replace the placeholder with "".

        `_build_process` does `f.get("duration", "unknown")` — a default that only
        applies when the key is ABSENT. An empty overlaid value would ship an
        empty ParameterValue, which is worse than the placeholder.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        chain = pipeline_mod._merge_plan_chain_names(
            {
                "process_chain": [
                    {"process_type": "Exposure", "name": "Dose", "parameters": {"duration": "  "}}
                ]
            }
        )

        assert "duration" not in self._hints_for(chain, "Exposure")

    def test_no_plan_still_yields_the_canonical_name_only_chain(self):
        """#262 regression guard — the no-provider crate stays byte-identical."""
        import builder.agents.pipeline.pipeline as pipeline_mod

        chain = pipeline_mod._merge_plan_chain_names(None)

        assert len(chain) == len(pipeline_mod._STANDARD_CHAIN)
        for step in chain:
            assert set(step["hints"]) == {"name"}


class TestConditionTableFromPlan:
    """#408 (b) — the plan's ``condition_table`` role must reach the populator.

    ``extract_plan`` classifies every plan file into
    ``["raw", "processed", "condition_table", "other"]`` and the pipeline threw the
    answer away: ``_attach_scanned_files`` re-derives a role from the filename via
    ``_file_role``, which only ever returns ``processed_data``/``raw_data``. So
    ``condition_table`` was unreachable by construction and every exported table
    shipped header-only while the per-well payload sat one directory away.
    """

    _PLATE = "plate_map.csv"
    _BODY = (
        "well_id,assay,cell_line,compound,concentration_value,concentration_unit\n"
        "A1,uptake,CHO-K1,Thyroxine,1.0,uM\n"
        "A2,uptake,CHO-K1,Thyroxine,10.0,uM\n"
    )

    def _state(self, tmp_path: Path, *, name: str | None = None, body: str | None = None):
        """A scanned state whose inventory holds a plate-map CSV under an approved root."""
        plate = tmp_path / (name or self._PLATE)
        plate.write_text(body if body is not None else self._BODY, encoding="utf-8")
        state = CrateState()
        state.metadata.title = "OATP1C1 uptake"
        state.metadata.description = "An in vitro uptake assay."
        state.approved_scan_roots.add(str(tmp_path.resolve()))
        state.scanned_files = [
            FileClassification(
                path=str(plate),
                filename=plate.name,
                size=plate.stat().st_size,
                mime_type="text/csv",
                first_rows=None,
            )
        ]
        return state, plate

    def _run(self, monkeypatch, state, plan_files, plan_extra=None):
        """Drive `_materialize_plan` with a stubbed leaf returning *plan_files*.

        The backbone is scaffolded first exactly as the spine's step 1 does, so the
        standard chain (#262) lays down the Exposure the table hangs off — without
        it there is no Exposure and the test would pass vacuously. ``plan_extra``
        merges further plan keys (e.g. ``compounds``) for tests that need the
        Exposure wired the way a real extraction would leave it.
        """
        import builder.agents.pipeline.pipeline as pipeline_mod

        plan = {
            "study": {"name": "OATP1C1 uptake", "description": "An uptake assay."},
            "files": plan_files,
        }
        plan.update(plan_extra or {})
        monkeypatch.setattr(pipeline_mod, "get_provider", lambda: "openai")
        monkeypatch.setattr(
            pipeline_mod,
            "extract_plan",
            lambda context, *, overrides=None, usage_sink=None: dict(plan),
        )
        engine = _engine(state)
        pipeline_mod._scaffold_backbone(engine)
        result = pipeline_mod._materialize_plan(engine)
        assert any(
            e.type == "LabProcess" and e.fields.get("process_type") == "Exposure"
            for e in engine.state.list_entities()
        ), "fixture is vacuous: no Exposure was scaffolded"
        return engine, result

    def _table_path(self, engine: AgentEngine) -> Path | None:
        from builder.tools._crate_mapping import _condition_table_rel, _mint_id

        for e in engine.state.list_entities():
            if e.type == "LabProcess" and e.fields.get("process_type") == "Exposure":
                out = engine.state.metadata.output_path
                if not out:
                    return None
                return Path(out) / _condition_table_rel(_mint_id(e))
        return None

    def test_a_basename_match_populates_the_table(self, monkeypatch, tmp_path: Path) -> None:
        """The core wiring: role=condition_table → real rows in the Exposure's CSV."""
        state, _ = self._state(tmp_path)
        state.metadata.output_path = str(tmp_path / "crate")
        engine, result = self._run(
            monkeypatch, state, [{"path": self._PLATE, "role": "condition_table"}]
        )

        table = self._table_path(engine)
        assert table is not None and table.is_file(), "no condition table was written"
        rows = table.read_text(encoding="utf-8").strip().splitlines()
        assert len(rows) > 1, "the table is still header-only — the role never reached the tool"
        assert "Thyroxine" in "\n".join(rows)
        assert result.get("condition_table")

    def test_the_model_can_only_return_a_basename_so_an_exact_path_would_never_fire(
        self, monkeypatch, tmp_path: Path
    ) -> None:
        """The silent-no-op guard (#408).

        ``_gather_context`` shows the leaf ``f.filename`` and never ``f.path``, so a
        plan path is always a bare basename. Matching on full-path equality would
        look wired and never fire once — this pins the basename contract by proving
        the scanned path and the plan path are NOT equal, yet the table populates.
        """
        state, plate = self._state(tmp_path)
        state.metadata.output_path = str(tmp_path / "crate")
        assert str(plate) != self._PLATE, "fixture must exercise basename != full path"

        engine, _ = self._run(
            monkeypatch, state, [{"path": self._PLATE, "role": "condition_table"}]
        )
        table = self._table_path(engine)
        assert table is not None and table.is_file()
        assert len(table.read_text(encoding="utf-8").strip().splitlines()) > 1

    def test_a_file_outside_the_approved_roots_is_refused(
        self, monkeypatch, tmp_path: Path
    ) -> None:
        """Fail-closed: plan paths are LLM free text and must not widen file access."""
        state, _ = self._state(tmp_path)
        state.metadata.output_path = str(tmp_path / "crate")
        # The inventory entry survives, but no approved root now contains it.
        state.approved_scan_roots.clear()
        state.approved_scan_roots.add(str((tmp_path / "elsewhere").resolve()))

        engine, result = self._run(
            monkeypatch, state, [{"path": self._PLATE, "role": "condition_table"}]
        )
        table = self._table_path(engine)
        assert (
            table is None
            or not table.is_file()
            or (len(table.read_text(encoding="utf-8").strip().splitlines()) == 1)
        ), "a file outside the approved roots must never be read"
        assert not (result.get("condition_table") or {}).get("populated")

    def test_two_candidates_refuse_to_guess(self, monkeypatch, tmp_path: Path) -> None:
        state, _ = self._state(tmp_path)
        second = tmp_path / "other_map.csv"
        second.write_text(self._BODY, encoding="utf-8")
        state.scanned_files.append(
            FileClassification(
                path=str(second),
                filename=second.name,
                size=second.stat().st_size,
                mime_type="text/csv",
                first_rows=None,
            )
        )
        state.metadata.output_path = str(tmp_path / "crate")

        engine, result = self._run(
            monkeypatch,
            state,
            [
                {"path": self._PLATE, "role": "condition_table"},
                {"path": "other_map.csv", "role": "condition_table"},
            ],
        )
        outcome = result.get("condition_table") or {}
        assert not outcome.get("populated"), "two candidates — the spine must not guess"
        assert outcome.get("reason"), "the refusal must be recorded, not silent"

    def test_no_condition_table_role_records_a_reason(self, monkeypatch, tmp_path: Path) -> None:
        state, _ = self._state(tmp_path)
        state.metadata.output_path = str(tmp_path / "crate")
        engine, result = self._run(monkeypatch, state, [{"path": self._PLATE, "role": "raw"}])
        outcome = result.get("condition_table") or {}
        assert not outcome.get("populated")
        assert outcome.get("reason")

    # -- #422: an unusable candidate must not beat having no candidate at all --

    def _offline_compound_seams(self, monkeypatch) -> None:
        """Canned offline lookup hit: a MISS mints no MolecularEntity at all, and
        the proposal needs the compound wired to the Exposure to have anything
        to propose. The CAS is the canned lookup value, never plan-fabricated (D5)."""
        import builder.tools.composites as composites_mod

        monkeypatch.setattr(
            composites_mod,
            "lookup_compound",
            lambda name: {
                "found": True,
                "data": {"cas": "51-48-9", "pubchem_cid": 5819, "source": "pubchem"},
                "error": None,
            },
        )

        def _fake_verify(state, entity_id, field):
            ent = state.get_entity(entity_id)
            if ent is not None:
                ent.set_field_status(field, "verified", "lookup")
            return {"verified": True, "entity_id": entity_id, "field": field}

        monkeypatch.setattr(composites_mod, "verify_identifier", _fake_verify)

    _COMPOUND_PLAN = {"compounds": [{"name": "Thyroxine", "role": "substrate"}]}

    def test_a_declined_candidate_falls_back_to_the_proposal(
        self, monkeypatch, tmp_path: Path
    ) -> None:
        """#422 — the S-VHPS26 shape in miniature: the plan labels a
        Parameter|Value metadata sheet ``condition_table``; the tool reads it and
        correctly refuses (no canonical column maps). Refusing must then fall
        back to the #438 proposal — a readable-but-unmappable candidate must not
        leave the crate with LESS than no candidate at all."""
        self._offline_compound_seams(monkeypatch)
        state, _ = self._state(
            tmp_path,
            name="assay_metadata.csv",
            body="Parameter,Value\nExposure duration_1,30\nTime unit_1,minutes\n",
        )
        state.metadata.output_path = str(tmp_path / "crate")
        engine, result = self._run(
            monkeypatch,
            state,
            [{"path": "assay_metadata.csv", "role": "condition_table"}],
            plan_extra=dict(self._COMPOUND_PLAN),
        )
        outcome = result.get("condition_table") or {}
        assert outcome.get("populated") is True, f"no fallback fired: {outcome}"
        assert outcome.get("proposed") is True
        assert outcome.get("fallback_from"), "the original refusal must be preserved"
        table = self._table_path(engine)
        assert table is not None and table.is_file()
        assert "Thyroxine" in table.read_text(encoding="utf-8")

    def test_an_unreadable_candidate_falls_back_to_the_proposal(
        self, monkeypatch, tmp_path: Path
    ) -> None:
        """#422 — a candidate no reader can open (here: a text file wearing an
        ``.xlsx`` suffix) must likewise fall back rather than ship header-only."""
        self._offline_compound_seams(monkeypatch)
        state, _ = self._state(tmp_path, name="plate_map.xlsx", body="this is not a workbook")
        state.metadata.output_path = str(tmp_path / "crate")
        engine, result = self._run(
            monkeypatch,
            state,
            [{"path": "plate_map.xlsx", "role": "condition_table"}],
            plan_extra=dict(self._COMPOUND_PLAN),
        )
        outcome = result.get("condition_table") or {}
        assert outcome.get("populated") is True, f"no fallback fired: {outcome}"
        assert outcome.get("proposed") is True
        assert outcome.get("fallback_from")
        table = self._table_path(engine)
        assert table is not None and table.is_file()
        assert "Thyroxine" in table.read_text(encoding="utf-8")

    def test_when_the_proposal_also_fails_the_primary_reason_survives(
        self, monkeypatch, tmp_path: Path
    ) -> None:
        """#422 — with nothing to propose (no compounds wired), the recorded
        reason must stay the PRIMARY failure (the decline), with the proposal's
        own failure alongside it — never silently swapped."""
        state, _ = self._state(
            tmp_path, name="assay_metadata.csv", body="Parameter,Value\nfoo,bar\n"
        )
        state.metadata.output_path = str(tmp_path / "crate")
        engine, result = self._run(
            monkeypatch,
            state,
            [{"path": "assay_metadata.csv", "role": "condition_table"}],
        )
        outcome = result.get("condition_table") or {}
        assert not outcome.get("populated")
        assert "column" in str(outcome.get("reason") or "").lower(), (
            f"the primary decline must survive as the reason: {outcome}"
        )
        assert outcome.get("proposal_reason"), (
            "the proposal's own failure must be recorded alongside"
        )
